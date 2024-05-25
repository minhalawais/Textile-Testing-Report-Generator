from create_test import create_test
from docx.enum.section import WD_SECTION
from docx.shared import Inches,Pt
from global_func import set_col_widths

def Total_test_template(doc,test_name,test_method,remarks,values,req):
    create_test(doc,test_name,test_method,remarks)
    # Define the maximum number of dictionary columns per table
    max_columns_per_table = 6

    num_tables = (len(values) + max_columns_per_table - 1) // max_columns_per_table



    for table_index in range(num_tables):
        start_column = table_index * max_columns_per_table
        end_column = start_column + max_columns_per_table

        num_columns = 1 + min(max_columns_per_table, len(values) - start_column)
        table = doc.add_table(rows=2, cols=num_columns+2)
        table.style = 'Table Grid'
        table.cell(0,0).text = "Substance Name"
        table.cell(1, 0).text = 'Total Lead'
        table.cell(0, 1).text = "CAS No"
        table.cell(1, 1).text = '50-00-0'

        for column_index, column_title in enumerate(list(values.keys())[start_column:end_column], start=2):
            table.cell(0, column_index).text = column_title

            row = table.rows[0]  # Start from the second row since the first row contains the headers
            row1 = table.rows[1]
            for column_index, column_values in enumerate(list(values.values())[start_column:end_column], start=2):
                cell = row.cells[column_index]
                cell.text = list(values.keys())[start_column+(column_index-2)]
                cell = row1.cells[column_index]
                cell.text = column_values
        a = table.cell(0,num_columns+1)
        a.text = "mg/kg"
        b = table.cell(1,num_columns+1)
        b.text = req[0]
        widths = (Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5),Inches(1.5), Inches(1.5), Inches(1.5))
        set_col_widths(table,widths)
        sample_heading = doc.add_paragraph()
        sample_heading.paragraph_format.space_before = Pt(0)
        sample_heading.paragraph_format.space_after = Pt(0)
    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'
    table.cell(0,0).text ="Note"
    table.cell(1,0).text = "MDL: 10 mg/kg"

    widths = (Inches(10), Inches(2))
    set_col_widths(table, widths)

def Total_test_data():

    values = {
        'A1+A2+A3': 'ND',
        'B1+B2+B3': 'ND',
        'C1+C2+C3': 'ND',
        'D1+D2+D3': 'ND',
        'E1+E4+E5': 'ND',
        'E2+E3': 'ND',
        'F1+F2+F3': 'ND',
        'F4+F5': 'ND',
        'G1+G3+G5': 'ND',
        'G2+G4': 'ND',
        'H1+H2+H3': 'ND'
    }
    test_name = "Total Lead (Pb) Content:"
    test_method = "DIN EN 16711-1:2016, Analyzed by ICP-OES"
    remarks = ""
    return test_name,test_method,remarks,values

