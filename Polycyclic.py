from create_test import create_test
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from global_func import set_col_widths

def Poly_test_template(doc,test_name,test_method,remarks,title,cas_no,values,req):
    create_test(doc,test_name,test_method,remarks)
    # Define the maximum number of dictionary columns per table
    max_columns_per_table = 5
    last = ['0.2','0.2','0.1','0.5','0.5','0.5','0.5','0.5','0.5','0.02','0.5','0.5']
    num_tables = (len(values) + max_columns_per_table) // max_columns_per_table



    for table_index in range(num_tables):


        start_column = table_index * max_columns_per_table
        end_column = start_column + max_columns_per_table

        num_columns = 2 + min(max_columns_per_table, len(values) - start_column)
        table = doc.add_table(rows=len(title)+1, cols=num_columns+1)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Substance Name'
        table.cell(0, 1).text = 'CAS No.'
        table.cell(0,num_columns).text = "Requirement"

        for column_index, column_title in enumerate(list(values.keys())[start_column:end_column], start=2):
            table.cell(0, column_index).text = column_title

        for i, (value, cas) in enumerate(zip(title, cas_no)):

            row = table.rows[i + 1]  # Start from the second row since the first row contains the headers

            cell1 = row.cells[0]
            cell1.text = value

            cell2 = row.cells[1]
            cell2.text = cas

            for column_index, column_values in enumerate(list(values.values())[start_column:end_column], start=2):
                cell = row.cells[column_index]
                if i < 25:
                    cell.text = column_values[i]

        sample_heading = doc.add_paragraph()
        widths = (Inches(1.5),  Inches(1), Inches(0.5), Inches(0.5),Inches(0.5),Inches(0.5),Inches(0.5),Inches(0.5),Inches(3))
        set_col_widths(table, widths)
    table = doc.add_table(rows=4, cols=1)
    table.style = 'Table Grid'
    table.cell(0,0).text ="Note"
    table.cell(1,0).text = "MDL: 0.2 mg/kg"
    table.cell(2,0).text = "Method Detection Limit"
    table.cell(3,0).text = "mg/kg: milligram per kilogram"
    widths = (Inches(10), Inches(1))
    set_col_widths(table, widths)
    doc.add_section(WD_SECTION.NEW_PAGE)
    create_test(doc, test_name, test_method, remarks)
def Poly_test_data():
    title = ["Benzo (a) anthracene", "Benzo (a) pyrene", "Benzo (b) fluoranthene", "Benzo [e] pyrene", "Benzo [j] fluoranthene", "Benzo (k) fluoranthene", "Chrysene", "Dibenzo(a,h)anthracene", "Naphthalene", "Acenaphthylene", "Acenaphtene", "Fluorene", "Phenanthrene", "Anthracene", "Fluoranthene", "Pyrene", "Indeno(1,2,3-cd) pyrene", "Benzo(g,h,i) perylene", "1-Methylpyrene", "Dibenzo[a,l]pyrene", "Dibenzo[a,i]pyrene", "Dibenzo[a,h]pyrene", "Dibenzo[a,e]pyrene", "Cyclopenta[c,d]pyrene", "Sum 24 PAHs", "Rating"]
    cas_no = [
        "56-55-3", "50-32-8", "205-99-2", "192-97-2", "205-82-3", "207-08-9", "218-01-9",
        "53-70-3", "91-20-3", "208-96-8", "83-32-9", "86-73-7", "85-01-8", "120-12-7",
        "206-44-0", "129-00-0", "193-39-5", "191-24-2", "2381-21-7", "191-30-0", "189-55-9",
        "189-64-0", "192-65-4", "27208-37-3", "-", "-"
    ]

    values = {
        'Sample A': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample B': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample C': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample D': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample E': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample F': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample G': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample H': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample I': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample J': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND'],
        'Sample K': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND','ND']
    }
    test_name = "Polycyclic Aromatic Hydrocarbon(PAH):"
    test_method = "AfPS GS 2019:01"
    remarks = ""
    return test_name,test_method,remarks,title,cas_no,values