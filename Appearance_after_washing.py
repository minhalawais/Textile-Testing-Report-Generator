from docx.shared import Pt,RGBColor,Inches
from create_test import create_test
from global_func import set_col_widths
def appearance_test_template(doc,test_name,test_method,remarks,title,req,sample_results):

    create_test(doc, test_name, test_method, remarks)
    for sample,name in sample_results.items():
        sample_heading = doc.add_heading()
        sample_run = sample_heading.add_run(f"{sample}")
        sample_run.font.name = 'Calibri'
        sample_run.font.size = Pt(12)
        sample_run.font.bold = True
        sample_run.font.underline = True
        sample_run.font.color.rgb = RGBColor(0, 0, 0)
        # Create a table for the sample
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        headers = ['Assesment', 'Result']
        header_cells = table.rows[0].cells
        for j, header in enumerate(headers):
            #header_cells[j].text = header
            row = table.rows[0]
            Nombre_text_formatted = row.cells[j].paragraphs[0].add_run(header)
            Nombre_text_formatted.bold = True
            Nombre_text_formatted.font.size = Pt(12)
        print(sample,name)
        for j in range(5):
            row_cells = table.add_row().cells
            row_cells[0].text = title[j]
            row_cells[1].text = sample_results[sample][title[j]]
            row_cells[2].text = req[j]
        widths = (Inches(2.5), Inches(2.5), Inches(2.5))
        set_col_widths(table, widths)


    doc.save('test_template.docx')




def appearance_test_data():
    sample_results = []
    test_name = "Appearance after washing: "
    test_method  = "DIN EN ISO 6330:2021/ 5077:2008"
    remarks = "Washing & drying procedure: Same as Dimensional stability to washing"
    title = ['Colour Change','Cross Staining','Appearance','Pilling / Fuzzing','Other Changes Observed']
    req = ['Class 3-4','Class 4-5','Satisfactory','Class 4-5','Not Accepted']
    sample_results = {'Sample A':{'Colour Change': '4-5','Cross Staining': '5','Appearance': 'No seam open\nNo stitch broken','Pilling / Fuzzing': 'Slight pilling class 4-5','Other Changes Observed': 'No other changes observed'},
                      'Sample B':{'Colour Change': '4-5','Cross Staining': '5','Appearance': 'No seam open\nNo stitch broken','Pilling / Fuzzing': 'Slight pilling class 4-5','Other Changes Observed': 'No other changes observed'}
                      ,'Sample C':{'Colour Change': '4-5','Cross Staining': '5','Appearance': 'No seam open\nNo stitch broken','Pilling / Fuzzing': 'Slight pilling class 4-5','Other Changes Observed': 'No other changes observed'}}
    return test_name,test_method,remarks,title,req,sample_results
