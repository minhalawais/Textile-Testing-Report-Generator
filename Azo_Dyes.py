from create_test import create_test
from docx.enum.section import WD_SECTION
from docx.shared import Inches,Pt
from global_func import set_col_widths

def Azo_test_template(doc,test_name,test_method,remarks,values,req):
    create_test(doc,test_name,test_method,remarks)
    # Define the maximum number of dictionary columns per table
    max_columns_per_table = 6

    num_tables = (len(values) + max_columns_per_table - 1) // max_columns_per_table



    for table_index in range(num_tables):
        start_column = table_index * max_columns_per_table
        end_column = start_column + max_columns_per_table

        num_columns = 1 + min(max_columns_per_table, len(values) - start_column)
        table = doc.add_table(rows=2, cols=num_columns+1)
        table.style = 'Table Grid'
        table.cell(1, 0).text = 'Result'

        for column_index, column_title in enumerate(list(values.keys())[start_column:end_column], start=1):
            table.cell(0, column_index).text = column_title

            row = table.rows[0]  # Start from the second row since the first row contains the headers
            row1 = table.rows[1]
            for column_index, column_values in enumerate(list(values.values())[start_column:end_column], start=1):
                cell = row.cells[column_index]
                cell.text = list(values.keys())[start_column+(column_index-1)]
                cell = row1.cells[column_index]
                cell.text = column_values
        a = table.cell(0,num_columns)
        a.text = "mg/kg"
        b = table.cell(1,num_columns)
        b.text = req[0]
        widths = (Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5),Inches(1.5), Inches(1.5), Inches(1.5))
        set_col_widths(table,widths)
        sample_heading = doc.add_paragraph()
        sample_heading.paragraph_format.space_before = Pt(0)
        sample_heading.paragraph_format.space_after = Pt(0)
    table = doc.add_table(rows=5, cols=1)
    table.style = 'Table Grid'
    table.cell(0,0).text ="Note"
    table.cell(1,0).text = "n.d. = not detected"
    table.cell(2,0).text = "mg/kg = ppm"
    table.cell(3,0).text = "* = Exceeds the limit	"
    table.cell(4,0).text = "Detection Limit = 5 mg/kg (for individual compound)	"
    widths = (Inches(10), Inches(2))
    set_col_widths(table, widths)
    data = [
        (1, "Biphenyl-4-ylamine\n4-aminobiphenyl xenylamine", "92-67-1", 14, "6-Methoxy-m-toluidine\np-Cresidine", "120-71-8"),
        (2, "Benzidine", "92-87-5", 15, "4,4'-Methylene-bis-(2-chloro-aniline)\n2,2'-Dichloro-4,4'-methylene-dianiline", "101-14-4"),
        (3, "4-Chlor-o-toluidine", "95-69-2", 16, "4,4'-Oxydianiline", "101-80-4"),
        (4, "2-Naphthylamine", "91-59-8", 17, "4,4'-Thiodianiline", "139-65-1"),
        (5, "o-Aminoazotoluene", "97-56-3", 18, "o-Toluidine 2-Aminotoluene", "95-53-4"),
        (6, "5-Nitro-o-toluidine\n4-Amino-2’,\n3-dimethylazobenzene\n4-o-Tolylazo-o-toluidine", "99-55-8", 19, "4-Methyl-m-phenylenediamine", "95-80-7"),
        (7, "4-Chloroaniline", "106-47-8", 20, "2,4,5-Trimethylaniline", "137-17-7"),
        (8, "4-Methoxy-m-phenylenediamine", "615-05-4", 21, "o-Anisidine\n2-Methoxyaniline", "90-04-0"),
        (9, "4,4’-Methylenedianiline\n4,4’-Diaminodiphenylmethane", "101-77-9", 22, "4-Amino azobenzene", "60-09-3"),
        (10, "3,3’-Dichlorobenzidine\n3,3’-Dichlorobiphenyl-4,\n4’-ylenediamine", "91-94-1", 23, "2,4-Xylidine", "95-68-1"),
        (11, "3,3'-Dimethoxybenzidine\no-Dianisidine", "119-90-4", 24, "2,6-Xylidine", "87-62-7"),
        (12, "3,3'-Dimethylbenzidine\n4,4'-Bi-o-toluidine", "119-93-7", 25, "Aniline", "62-53-3"),
        (13, "4,4'-Methylenedi-o-toluidine", "838-88-0", 26, "4-Aminoaniline\n1,4-Phenylenediamine", "106-50-3"),
    ]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    runner = p.add_run("List of Azo Dyes:")
    runner.bold = True
    runner.italic = True

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Set table header and bold the first row
    headings = ['Sr#', 'Substance name', 'CAS No.', 'Sr#', 'Substance name', 'CAS No.']
    for i, heading in enumerate(headings):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        paragraph.text = heading
        run = paragraph.runs[0]
        run.font.bold = True

    # Insert data into the table
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = str(cell_value)

    # Adjust column widths
    for col in table.columns:
        for cell in col.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size (optional)
    widths = (Inches(0.3), Inches(2.5), Inches(1.2), Inches(0.3), Inches(2.5), Inches(1.2))
    set_col_widths(table, widths)

def Azo_test_data():

    values = {
        'A1+A2+A3': 'ND',
        'B1+B2+B3': 'ND',
        'C1+C2+C3': 'ND',
        'D1+D2+D3': 'ND',
        'E1+E4+E5': 'ND',
        'E2+E3': 'ND',
        'F1+F2+F3': 'ND',
        'F4+F5': 'ND',
        'G1+G3+G5': 'ND',
        'G2+G4': 'ND',
        'H1+H2+H3': 'ND'
    }
    test_name = "Azo-Dyes (including Aniline):"
    test_method = "All Textile: According to DIN EN ISO 14362-1:2017 – Analysis was conducted with GC-MS/HPLC-DAD. "
    remarks = "Determination of 4-aminoazobenzene (CAS No.:60-09-3) –DIN EN ISO 14362-3:2017; with the use of Gas Chromatography – Mass Spectrometry (GC-MS)"
    return test_name,test_method,remarks,values

