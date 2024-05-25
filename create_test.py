from docx.shared import Pt,RGBColor,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from global_func import set_col_widths
def create_test(doc,test_name,test_method,remarks):
    table = doc.add_table(rows=1, cols=2)

    # Add the test_name heading in column 1 (cell 0, 0)
    cell_0_0 = table.cell(0, 0)
    paragraph_0_0 = cell_0_0.paragraphs[0]

    run_0_0 = paragraph_0_0.add_run(test_name)
    run_0_0.font.name = 'Calibri'
    run_0_0.font.size = Pt(12)
    run_0_0.font.bold = True
    run_0_0.font.underline = True

    paragraph_0_0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add "requirements" in the second cell (cell 0, 1)
    requirements_text = "Requirements\t"
    cell_0_1 = table.cell(0, 1)
    paragraph_0_1 = cell_0_1.paragraphs[0]
    run_0_1 = paragraph_0_1.add_run(requirements_text)
    run_0_1.font.name = 'Calibri'
    run_0_1.font.size = Pt(10)
    run_0_1.font.bold = True
    run_0_1.font.underline = True
    paragraph_0_1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    widths = (Inches(4), Inches(5))
    set_col_widths(table, widths)
    # Add the test method
    method_paragraph = doc.add_paragraph()
    method_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    method_paragraph.add_run(test_method)
    method_paragraph.paragraph_format.space_before = Pt(0)
    method_paragraph.paragraph_format.space_after = Pt(0)
    remarks_paragraph = doc.add_paragraph()
    remarks_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    remarks_paragraph.add_run(remarks)
    remarks_paragraph.paragraph_format.space_before = Pt(0)
    remarks_paragraph.paragraph_format.space_after = Pt(0)
