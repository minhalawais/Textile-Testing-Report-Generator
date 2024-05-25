from docx.enum.section import WD_SECTION
from docx.shared import Pt,RGBColor,Inches
from create_test import create_test
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from global_func import set_col_widths




def fiber_test_template(doc,test_name,test_method,remarks,title,sample_results,req):
    create_test(doc,test_name,test_method,remarks)
    for i,sample in enumerate(sample_results):
        if i % 3 == 0 and i !=0:
            doc.add_section(WD_SECTION.NEW_PAGE)
            create_test(doc, test_name, test_method, remarks)

        temp = list(sample.keys())
        temp = temp[0]

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        headers = [temp, 'Labeled Fiber Content', 'Actual (Tested) Fiber Content', 'Suggested Fiber Content']
        for j, header in enumerate(headers):
            row = table.rows[0]
            Nombre_text_formatted = row.cells[j].paragraphs[0].add_run(header)
            Nombre_text_formatted.underline = True
            Nombre_text_formatted.font.size = Pt(10)

        for i in range(4):
            row_cells = table.add_row().cells
            row_cells[0].text = title[i]
            if i == 0:
                for j in range(1,4):
                    row_cells[j].text = sample[temp][0]['Cotton'][j-1]
            elif i==1:
                for j in range(1, 4):
                    row_cells[j].text = sample[temp][0]['Polyamide'][j-1]
            elif i==2:
                for j in range(1, 4):
                    row_cells[j].text = sample[temp][0]['Polyester'][j-1]
            else:
                for j in range(1, 4):
                    row_cells[j].text = sample[temp][0]['Elastane'][j-1]
        a = table.cell(0,4)
        for i in range(1,5):
            b = table.cell(i,4)
            A = a.merge(b)
            a = A
        a.text = req[0]
        a.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        doc.add_paragraph()
        widths = (Inches(2), Inches(1.3), Inches(1.3), Inches(1.3), Inches(1.5))
        set_col_widths(table, widths)

def fiber_test_data():
    title = ['Cotton (%)','Polyamide (%)','Polyester (%)','Elastane (%)']
    test_name = "Fiber Composition:"
    test_method = "DIN EN ISO 1833-7:2017,11:2017"
    remarks = ''
    values = [{'Sample A': [
        {'Cotton': ['72', '69.2', '69'], 'Polyamide': ['72', '69.2', '69'], 'Polyester': ['72', '69.2', '69'],
         'Elastane': ['72', '69.2', '28']}]},
              {'Sample B': [
                  {'Cotton': ['72', '69.2', '69'], 'Polyamide': ['72', '69.2', '69'], 'Polyester': ['72', '69.2', '69'],
                   'Elastane': ['72', '69.2', '28']}]},
              {'Sample C': [
                  {'Cotton': ['72', '69.2', '69'], 'Polyamide': ['72', '69.2', '69'], 'Polyester': ['72', '69.2', '69'],
                   'Elastane': ['72', '69.2', '28']}]},
              {'Sample D': [
                  {'Cotton': ['72', '69.2', '69'], 'Polyamide': ['72', '69.2', '69'], 'Polyester': ['72', '69.2', '69'],
                   'Elastane': ['72', '69.2', '28']}]}]

    return test_name,test_method,remarks,title,values

"""
test_name = 'Dimensional stability to washing'
test_method = 'DIN EN ISO 6330:2021/ 5077:2008'
remarks = 'Machine wash at (30°c) in household washing machine with Persil detergent, Normal cycle, 2.0 kg wash load, Flat dry.'
values = [{'Sample A':[{'Cotton':['72','69.2','69'],'Polyamide':['72','69.2','69'],'Polyester':['72','69.2','69'],'Elastane':['72','69.2','28']}]},
          {'Sample B':[{'Cotton':['72','69.2','69'],'Polyamide':['72','69.2','69'],'Polyester':['72','69.2','69'],'Elastane':['72','69.2','28']}]},
          {'Sample C':[{'Cotton':['72','69.2','69'],'Polyamide':['72','69.2','69'],'Polyester':['72','69.2','69'],'Elastane':['72','69.2','28']}]},
          {'Sample D':[{'Cotton':['72','69.2','69'],'Polyamide':['72','69.2','69'],'Polyester':['72','69.2','69'],'Elastane':['72','69.2','28']}]}]
test_name,test_method,remarks,title,sample_results = dimensional_test_data(test_name,test_method,remarks,values)
dimensional_test_template(doc,test_name,test_method,remarks,title,sample_results)
"""