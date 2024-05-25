from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from generate_qrcode import generate_qr_code_with_image
from global_func import requirement_data,doc_to_jpg,add_image_to_docx,convert_to_pdf,mainPageRemainingData,insert_image_in_docx_table,list_images_in_folder,set_secondarya_req,set_secondaryb_req,set_primary_req,count_sections,main,get_image_links,create_temp_document,add_row_to_table,add_care_labels,add_arrow,add_arrow_to_word
from flask import Flask,request,render_template,jsonify,redirect,url_for
import asyncio
import os
import base64

app = Flask(__name__)
app.jinja_env.auto_reload = True
app.config['TEMPLATES_AUTO_RELOAD'] = True
# Initialize variables to control the test flow
global current_test_index
current_test_index = 0
test_data = []  # Fill this with your test data

# Use an asyncio Event to signal when to continue
continue_event = asyncio.Event()
app = Flask(__name__)
temp_document =Document()

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)
font_name = 'Calibri'
font_size = Pt(10)


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


physical_tests = {'P': ['Colour fastness to rubbing','Colour fastness to water','Dimensional stability to washing','Seam spirality after laundering','Appearance after washing: '
                        'Fabric weight','Fiber Composition'],'P 1':['Colorfastness to saliva and perspiration','Tensile test on small parts']}

chemical_tests = {'C1':['Azo-Dyes','Aromatic Amine Salts','Alkylphenole/ Alkylphenolethoxylate (AP/APEO)','Quinoline','Disperse Dyes','Extractable (heavy) Metals',
                        'Formaldehyde'],'C2':['Azo-Dyes','Aromatic Amine Salts','Alkylphenole/ Alkylphenolethoxylate (AP/APEO)','Quinoline','Disperse Dyes','Extractable (heavy) Metals',
                        'Formaldehyde','Nickel release','Total Lead (Pb) Content','Total Cadmium (Cd) Content'],'C3':['Azo-Dyes','Aromatic Amine Salts','Alkylphenole/ Alkylphenolethoxylate (AP/APEO)','Quinoline','Disperse Dyes','Extractable (heavy) Metals',
                        'Formaldehyde','Nickel release','Total Lead (Pb) Content','Total Cadmium (Cd) Content','Phthalates','Chlorinated Paraffin’s (SCCP & MCCP)','Polycyclic Aromatic Hydrocarbon (PAH)']}
special_package = {'PC A':['Colour fastness to rubbing','Dimensional stability to washing','Colour fastness to water','Colour fastness to perspiration','Appearance after washing','Fiber Composition',
                           'Azo-Dyes','Aromatic Amine Salts','Alkylphenole/ Alkylphenolethoxylate (AP/APEO)'],'PC B':['Colour fastness to rubbing','Dimensional stability to washing','Colour fastness to water','Colour fastness to perspiration','Appearance after washing','Fiber Composition',
                           'Azo-Dyes','Aromatic Amine Salts','Alkylphenole/ Alkylphenolethoxylate (AP/APEO)','Colorfastness to saliva and perspiration','Tensile test on small parts'],'BSL':['Colour fastness to rubbing','Dimensional stability to washing','Colour fastness to water''Seam spirality after laundering','Fabric weight','Appearance after washing','Fiber Composition',
                           'Azo-Dyes','Aromatic Amine Salts','Alkylphenole/ Alkylphenolethoxylate (AP/APEO)','Extractable (heavy) Metals','Quinoline','Disperse Dyes','Formaldehyde','Total Lead (Pb) Content','Total Cadmium (Cd) Content','Phthalates','Chlorinated Paraffin’s (SCCP & MCCP)','Polycyclic Aromatic Hydrocarbon (PAH)','Chromium'
                            ,'Colorfastness to saliva and perspiration','Tensile test on small parts','Migration of certain substances','Flame retardants','Thermal hazards','Neck openings'],'BS':['Azo-Dyes','Aromatic Amine Salts','Alkylphenole/ Alkylphenolethoxylate (AP/APEO)','Extractable (heavy) Metals','Quinoline','Disperse Dyes','Formaldehyde','Total Lead (Pb) Content','Total Cadmium (Cd) Content','Phthalates','Chlorinated Paraffin’s (SCCP & MCCP)','Polycyclic Aromatic Hydrocarbon (PAH)'
                            ,'Colorfastness to saliva and perspiration','Tensile test on small parts']}
sample_picking = {'RP':['Solvents','Dye Carrier, Solvents'],'RPS':['Solvents','Dye Carrier, Solvents','Bisphenol A','Extractable (heavy) Metals','Disperse Dyes','Quinoline','Formaldehyde']}
applicant_list = ['Applicant:', 'Contact:', 'Address:', 'Tel:', 'E-mail:']
applicant_dict = {'Applicant:': 'Kamal Mills Private Limited.', 'Contact:': 'Muhammad Asad',
                  'Address:': '3 KM, Jhumra Road, Khurrianwala, Faisalabad-Pakistan.', 'Tel:': '0300-8793742',
                  'E-mail:': 'Muhammad.asad@kamal.com.pk'}
buyer_list = ['Buyer Name:', 'Supplier Name:', 'Agent:', 'Country of Origin:', 'Country of Destination:']
buyer_dict = {'Buyer Name:': 'Kik Textilien und Non-Food GmbH', 'Supplier Name:': 'Kamal Mills Private Limited.',
              'Agent:': 'Matrix Sourcing', 'Country of Origin:': 'Pakistan', 'Country of Destination:': 'Germany'}
sample_list = ['Product Description:', 'Material Composition:', 'Fabric:', 'Fabric Weight:',
               'Merchandise Category (WGR):', 'P.O Reference No:', 'P.O:', 'Article No:', 'Article Description:',
               'Style Color:', 'Supplier No:', 'Season:', 'Buying Dept (EKB) :', 'Dye stuff:',
               'Previous Report# (for retest):']

@app.route('/')
def index():
    return render_template('loginPage.html')

@app.route('/SelectionScreen', methods=['GET'])
def selection_screen():
    global sample_dict, test_package
    search_query = request.args.get('req')
    remData = mainPageRemainingData()
    sample_dict = {'Product Description:': 'Children boy ergee socks 4 pair',
                   'Material Composition:': '  72% Cotton, 26% Polyamide, 2% Elastane', 'Fabric:': 'Jersey  ',
                   'Fabric Weight:': '/', 'Merchandise Category (WGR):': '837', 'P.O Reference No:': 'P203649',
                   'P.O:': remData['P.O:'], 'Article No:': remData['Article No:'],
                   'Article Description:': 'Children boy ergee socks 4 pair', 'Style Color:': 'Orange, Green',
                   'Supplier No:': remData['Supplier No:'], 'Season:': '123', 'Buying Dept (EKB) :': remData['Buying Dept (EKB) :'], 'Dye stuff:': '/',
                   'Previous Report# (for retest):': '/'}
    test_package = remData['Test Package:']
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    for i in range(5):
        row = table.rows[i]
        row_cells = row.cells[0].paragraphs[0].add_run(applicant_list[i])
        row_cells.bold = True
        row_cells.font.size = Pt(10)
        row_cells1 = row.cells[1].paragraphs[0].add_run(applicant_dict[applicant_list[i]])
        row_cells = row.cells[3].paragraphs[0].add_run(buyer_list[i])
        row_cells.bold = True
        row_cells.font.size = Pt(10)
        row_cells = row.cells[4].paragraphs[0].add_run(buyer_dict[buyer_list[i]])

    sample_heading = doc.add_paragraph()

    widths = (Inches(1), Inches(3), Inches(0.2), Inches(1.5), Inches(3))
    set_col_widths(table, widths)

    # second table
    table = doc.add_table(rows=16, cols=3)
    table.style = 'Table Grid'

    for i, item in enumerate(sample_list, start=1):
        cell1 = table.cell(i, 1)
        cell2 = table.cell(i, 2)

        row_cells = cell1.paragraphs[0].add_run(item)
        row_cells.bold = True
        row_cells.font.size = Pt(10)
        row_cells1 = cell2.paragraphs[0].add_run(sample_dict.get(item, ''))
    a = table.cell(1, 0)
    for j in range(1, 16):
        b = table.cell(j, 0)
        A = a.merge(b)
        a = A
    a = table.cell(0, 0)
    b = table.cell(0, 1)
    c = table.cell(0, 2)
    A = a.merge(b)
    B = A.merge(c)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w')))
    B._tc.get_or_add_tcPr().append(shading_elm_1)
    B.text = "Sample Information	"
    cell = table.cell(1, 0)
    paragraph = cell.paragraphs[0]

    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sample_heading = doc.add_paragraph()
    widths = (Inches(1.5), Inches(2), Inches(4))
    set_col_widths(table, widths)

    # third table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    widths = (Inches(3), Inches(8))
    set_col_widths(table, widths)
    cell1 = table.cell(0, 0)
    cell2 = table.cell(1, 0)
    row_cells = cell1.paragraphs[0].add_run("Submitted Care instructions:")
    row_cells.bold = True
    row_cells.font.size = Pt(10)
    row_cells = cell2.paragraphs[0].add_run("Test Package:	")
    row_cells.bold = True
    row_cells.font.size = Pt(10)
    cell2 = table.cell(1, 1)
    cell2.text = test_package
    # Generate QrCode
    text = f"Report# 02249-23/ Issued Date: 06 02 02/ Rating: Pass/ Applicant: {applicant_dict['Applicant:']}/ Client: {buyer_dict['Buyer Name:']}/ Sample: {sample_dict['Product Description:']}"
    image_path = r"Logo\tti_logo.png"

    generate_qr_code_with_image(text, image_path, )

    # fourth table
    sample_heading = doc.add_paragraph()
    sample_run = sample_heading.add_run("")
    table = doc.add_table(rows=1, cols=3)
    cell1 = table.cell(0, 0)
    paragraph = cell1.paragraphs[0]
    paragraph.alignment = 0
    run = paragraph.add_run()
    run.add_picture(r"Logo\qrcode.png", width=Inches(1.5))
    paragraph.alignment = 0
    cell3 = table.cell(0, 2)
    paragraph = cell3.paragraphs[0]
    sample_run = paragraph.add_run("For and on behalf of\nTEXTILE TESTING INTERNATIONAL")
    paragraph = cell3.add_paragraph()
    paragraph.alignment = 1
    sample_run = paragraph.add_run()
    sample_run.add_picture(r"Logo\signature.png", width=Inches(2))
    paragraph = cell3.add_paragraph()
    sample_run = paragraph.add_run("Ali Ashraf\n    AVP Softlines ")
    sample_run.bold = True
    paragraph.alignment = 1
    widths = (Inches(2.59), Inches(3), Inches(2.18))
    set_col_widths(table, widths)
    doc.add_section(WD_SECTION.NEW_PAGE)
    for row in table.rows:
        row.height = Inches(1.4)
    # test names page
    sample_run.font.size = Pt(8)

    from test_name_page import test_name_table
    test_name_table(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()
    sample_run = sample_heading.add_run("\n")

    # component breakdown page
    from component_breakdown import breakdown_table
    breakdown_table(doc)


    # collage image

    from images_collage import set_image
    doc.add_section(WD_SECTION.NEW_PAGE)
    table = doc.add_table(rows=1, cols=3)
    set_image(doc)


    return render_template('selectionScreen.html')

@app.route('/create_report', methods=['Post','GET'])
def create_report():
    data = request.get_json()
    option = data.get('option')
    if option == 'button-1':
        requirement = set_primary_req(sample_dict['Fabric:'], test_package, sample_dict['Style Color:'])
    elif option == 'button-2':
        requirement = set_secondarya_req(sample_dict['Fabric:'], test_package, sample_dict['Style Color:'])
    else:
        requirement = set_secondaryb_req(sample_dict['Fabric:'], test_package, sample_dict['Style Color:'])

    global current_test_index
    get_images = []
    # Receive input data from the client
    starting_index = 0
    current_test_index += 1
    input_data = request.get_json()
    input_data = list(input_data.values())
    # appearance after washing

    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()
    from Dimensional_stability_to_washing import dimensional_test_data, dimensional_test_template
    test_name, test_method, remarks, data = dimensional_test_data()
    dimensional_test_template(doc, test_name, test_method, remarks, data, requirement['Dimensional stability to washing'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()
    doc.save('test_template.docx')
    # doc_to_jpg(range(last_page,count_sections(doc)))
    # colour fastness to water
    from Appearance_after_washing import appearance_test_data, appearance_test_template
    test_name, test_method, remarks, title, req, sample_results = appearance_test_data()
    appearance_test_template(doc, test_name, test_method, remarks, title, requirement['Appearance after washing'], sample_results)
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()
    input_data_list = [['Change 3-4'],['/'],['Staining 3-4']]
    from Colour_fastness_to_water import color_test_data, color_test_template
    test_name, test_method, remarks, title, values = color_test_data()
    color_test_template(doc, test_name, test_method, remarks, title, values,requirement['Colour fastness to water'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()
    input_data_list = [['Change 3-4'],['Contrast staining 4-5'],['Staining 3-4']]
    # colour fastness to perspiration

    from Colour_fastness_to_perspiration import color_fp_test_data, color_fp_test_template
    test_name, test_method, remarks, title, values = color_fp_test_data()
    color_fp_test_template(doc, test_name, test_method, remarks, title, values,requirement['Colour fastness to perspiration'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()
    input_data_list = [['Dry 4'], ['Wet 2-3']]
    # colour fastness to water
    from Colour_fastness_to_rubbing import Color_fr_test_data, Color_fr_test_template
    test_name, test_method, remarks, title, values = Color_fr_test_data()
    Color_fr_test_template(doc, test_name, test_method, remarks, title, values,requirement['Colour fastness to rubbing:'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()
    sample_run = sample_heading.add_run("\n")
    input_data_list = [['±3%']]
    # fiber composition
    from fiber_composition import fiber_test_template, fiber_test_data
    test_name = "Fiber Composition:"
    test_method = "DIN EN ISO 1833-7:2017,11:2017"
    remarks = ''
    test_name, test_method, remarks, title, sample_results = fiber_test_data()
    fiber_test_template(doc, test_name, test_method, remarks, title, sample_results,requirement['Fiber Composition:'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['Sum of \nNP, OP: 5 mg/kg'], ['Sum of \nNPEO, OPEO: 50 mg/kg'],['/']]
    # Alkylphenole test
    from Alkylphenole import Alkylphenole_test_data, Alkylphenole_test_template
    test_name, test_method, remarks, title, cas_no, values = Alkylphenole_test_data()
    Alkylphenole_test_template(doc, test_name, test_method, remarks, title, cas_no, values,['Sum of NP, OP: 5 mg/kg','Sum of NPEO, OPEO: 50 mg/kg','/'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['±5%'], ['190 g/m²'], ['5.60 Oz']]
    # Fabric Weight
    from fabric_weight import fabric_test_data, fabric_test_template
    test_name, test_method, remarks, title, values = fabric_test_data()
    fabric_test_template(doc, test_name, test_method, remarks, title, values,requirement['Fabric weight:'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['±3%'], ['3 cm']]
    # Seam spirality after laundering:
    from Seam_spirality_after_laundering import seam_test_template, seam_test_data
    test_name, test_method, remarks, title, values = seam_test_data()
    seam_test_template(doc, test_name, test_method, remarks, title, values,requirement['Seam spirality after laundering:'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['20']]
    # Azo-Dyes (including Aniline):
    from Azo_Dyes import Azo_test_template, Azo_test_data
    test_name, test_method, remarks, values = Azo_test_data()
    Azo_test_template(doc, test_name, test_method, remarks, values,['20'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['20']]
    # Aromatic Amine Salts:
    from Aromatic_Amine_Salts import Aromatic_test_data, Aromatic_test_template
    test_name, test_method, remarks, values = Aromatic_test_data()
    Aromatic_test_template(doc, test_name, test_method, remarks, values,['20'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['ND']]
    from Formaldehyde import Formaldehyde_test_data, Formaldehyde_test_template
    test_name, test_method, remarks, values = Formaldehyde_test_data()
    Formaldehyde_test_template(doc, test_name, test_method, remarks, values,['ND'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['Individual Test: 500 \n2-1 composite: 225\n3-1 composite: 150','75','90']]
    # Total Lead
    from Total_Lead_Content import Total_test_data, Total_test_template
    test_name, test_method, remarks, values = Total_test_data()
    Total_test_template(doc, test_name, test_method, remarks, values,['75'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['Individual Test: 100 \n2-1 composite: 45\n3-1 composite: 30','40']]
    # Total Cadmium (Cd) Content:
    from Total_Cadmium import Cadmium_test_template, Cadmium_test_data
    test_name, test_method, remarks, values = Cadmium_test_data()
    Cadmium_test_template(doc, test_name, test_method, remarks, values,['40'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['0.2'],['0.2'],['0.1'],['0.5'],['30'],['1.0'],['1.0'],['25'],['1.0'],['0.02'],['1000'],['100']]
    # Extractable Metals
    from Extractable_metals import extract_test_data, extract_test_template
    test_name, test_method, remarks, title, cas_no, mdl, values = extract_test_data()
    extract_test_template(doc, test_name, test_method, remarks, title, cas_no, mdl, values, ['0.2','0.2','0.1','0.5','30','1.0','1.0','25','1.0','0.02','1000','100'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['Individual Test :50 \n2-1 composite: 22.5\n3-1 composite: 15','50']]
    # Quinoline
    from Quinoline import Quinoline_test_template, Quinoline_test_data
    test_name, test_method, remarks, values = Quinoline_test_data()
    Quinoline_test_template(doc, test_name, test_method, remarks, values,['Individual Test :50 \n2-1 composite: 22.5\n3-1 composite: 15','50'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    input_data_list = [['Individual Test :1000\n2-1 composite: 450\n3-1 composite: 300','Each :100 \nSum of all: 250']]
    # Phthalates
    from phthalates import Ph_test_data, Ph_test_template
    test_name, test_method, remarks, values = Ph_test_data()
    Ph_test_template(doc, test_name, test_method, remarks, values,['Individual Test :1000\n2-1 composite: 450\n3-1 composite: 300','Each :100 \nSum of all: 250'])
    input_data_list = [['8 PAHs \n(substance limit):\n0.5 mg/kg'],['2.0 mg/kg'],['no single limit'],['5.0 mg/kg']]

    doc.save("test_template.docx")
    # Polycyclic
    from Polycyclic import Poly_test_data, Poly_test_template
    test_name, test_method, remarks, title, cas_no, values = Poly_test_data()
    Poly_test_template(doc, test_name, test_method, remarks, title, cas_no, values,['8 PAHs \n(substance limit):\n0.5 mg/kg','2.0 mg/kg','no single limit','5.0 mg/kg'])
    doc.save("test_template.docx")

    from header import set_header
    set_header(doc)
    doc.save("test_template.docx")
    for i in range(5):
        import time
        time.sleep(1)
        try:
            from footer import set_footer
            set_footer()
            break
        except Exception as e:
            print("Error", e)
    return render_template('images.html')


@app.route('/add_labels')
def add_labels():
    print('Yes')
    return render_template('images.html')

@app.route('/save_images', methods=['POST'])
def save_images():
    selected_images = {}
    for section in range(1, 6):
        selected_images[f'section{section}'] = request.form.get(f'section{section}')
    print(selected_images)
    add_care_labels(doc,selected_images)

    return render_template('collageScreen.html',imageUrl="..\static\collage.png")
@app.route('/save_front_images', methods=['POST','GET'])
def save_front_images():

    return render_template('frontPageImage.html',imageUrl=r"..\static\mainImage.jpg")
@app.route('/send_label_and_position', methods=['POST'])
def send_label_and_position():
    data = request.json
    label = data.get('label')
    cursor_position = data.get('cursorPosition')
    if cursor_position == 'right':
        cursor_position = True
    else:
        cursor_position = False
    print(label,cursor_position)
    image_link = add_arrow(label,cursor_position)
    # Process the label and cursor position, and get the image link

    response_data = {
        "message": "Data received and processed successfully."
    }

    return jsonify(response_data),200

@app.route('/save_edited_image', methods=['POST'])
def save_edited_image():
    try:
        data = request.get_json()
        image_data = data['imageDataURL']

        # Extract image data from the data URL
        image_data = image_data.split(',')[1]  # Remove data:image/png;base64,
        image_data = base64.b64decode(image_data)

        # Define a path to save the image (adjust as needed)
        save_path = r'static'
        image_filename = os.path.join(save_path, 'edited_image.png')

        # Save the image to the specified path
        with open(image_filename, 'wb') as image_file:
            image_file.write(image_data)
        add_image_to_docx(doc)
        doc.save('test_template.docx')
        # Respond with a success message
        response = {'message': 'Image data received and saved successfully'}
        return jsonify(response), 200
    except Exception as e:
        response = {'message': 'Error saving image data'}
        return jsonify(response), 500
@app.route('/save_front_edited_image', methods=['POST'])
def save_front_edited_image():
    try:
        data = request.get_json()
        image_data = data['imageDataURL']

        # Extract image data from the data URL
        image_data = image_data.split(',')[1]  # Remove data:image/png;base64,
        image_data = base64.b64decode(image_data)

        # Define a path to save the image (adjust as needed)
        save_path = r'static'
        image_filename = os.path.join(save_path, 'edited_front_image.png')

        # Save the image to the specified path
        with open(image_filename, 'wb') as image_file:
            image_file.write(image_data)
        insert_image_in_docx_table(doc)
        doc.save('test_template.docx')
        # Respond with a success message
        response = {'message': 'Image data received and saved successfully'}
        return jsonify(response), 200
    except Exception as e:
        response = {'message': 'Error saving image data'}
        return jsonify(response), 500
@app.route('/test_result')
def test_result():
    names = ["Dimensional stability to washing", "Appearance after washing", "Colour fastness to water",
             "Colour fastness to perspiration", "Colour fastness to rubbing", "Fiber Composition",
             "Alkylphenole/ Alkylphenolethoxylate (AP/APEO)", "Fabric weight", "Seam spirality after laundering",
             "Azo-Dyes (including Aniline)", "Aromatic Amine Salts", "Formaldehyde", "Total Lead (Pb) Content",
             "Total Cadmium (Cd) Content", "Extractable (heavy) Metals", "Quinoline", "Phthalates",
             "Polycyclic Aromatic Hydrocarbon(PAH)"]
    return render_template('testResult.html', names=names)

@app.route('/showPdf', methods=['POST','GET'])
def showPdf():
    convert_to_pdf()
    return render_template('downloadPage.html',pdfUrl='../static/test_template.pdf')

@app.route('/submit_test_result', methods=['POST'])
def submit_test_result():
    data = request.json
    names = ["Dimensional stability to washing", "Appearance after washing", "Colour fastness to water",
             "Colour fastness to perspiration", "Colour fastness to rubbing", "Fiber Composition",
             "Alkylphenole/ Alkylphenolethoxylate (AP/APEO)", "Fabric weight", "Seam spirality after laundering",
             "Azo-Dyes (including Aniline)", "Aromatic Amine Salts", "Formaldehyde", "Total Lead (Pb) Content",
             "Total Cadmium (Cd) Content", "Extractable (heavy) Metals", "Quinoline", "Phthalates",
             "Polycyclic Aromatic Hydrocarbon(PAH)"]
    for i in range(len(names)):
        add_row_to_table(doc,i+1,names[i],data[names[i]])
        # Process the data as needed
    doc.save('test_template.docx')
    print(data)  # Example: Print the received data in the console
    return jsonify({'message': 'Data received successfully'})
if __name__ == '__main__':
    app.run(debug=True)