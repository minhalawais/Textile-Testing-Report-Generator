from docx.shared import Pt,RGBColor,Inches
from create_test import create_test
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def dimensional_test_template(doc,test_name,test_method,remarks,sample_results,requirement):
    create_test(doc,test_name,test_method,remarks)
    for sample,values in sample_results.items():
        sample_heading = doc.add_heading()
        sample_run = sample_heading.add_run(f"{sample}")
        sample_run.font.name = 'Calibri'
        sample_run.font.size = Pt(12)
        sample_run.font.bold = True
        sample_run.font.underline = True
        sample_run.font.color.rgb = RGBColor(0, 0, 0)
        # Create a table for the sample
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        headers = ['Points Of Measurement', 'Before Wash (Cm)', 'After Wash (Cm)', 'Dimensional Change (%)']
        para = ['Product Length','Product Width']
        header_cells = table.rows[0].cells
        for j, header in enumerate(headers):
            header_cells[j].text = header
        temp = values[0]
        for i in range(2):
            row_cells = table.add_row().cells
            row_cells[0].text = para[i]
            row_cells[1].text = temp['Before Wash'][i]
            row_cells[2].text = temp['After Wash'][i]
            row_cells[3].text = temp['Dimensional Change'][i]

        table.add_column(Inches(2))
        a = table.cell(0, 4)
        b = table.cell(1, 4)
        c = table.cell(2,4)
        A = a.merge(b)
        B = A.merge(c)
        B.text = requirement
        B.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    remarks_paragraph = doc.add_paragraph()
    remarks_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    remarks_paragraph.add_run("(+) Denotes Extension  (-) Shrinkage")
    remarks_paragraph.paragraph_format.space_before = Pt(15)


def dimensional_test_data():
    test_name = 'Dimensional stability to washing'
    test_method = 'DIN EN ISO 6330:2021/ 5077:2008'
    remarks = 'Machine wash at (30°c) in household washing machine with Persil detergent, Normal cycle, 2.0 kg wash load, Flat dry.'

    data = {'Sample A:': [
        {'Before Wash': ['18.2', '18.8'], 'After Wash': ['17.8', '15.5'], 'Dimensional Change': ['-0.55', '0.22']}],
            'Sample B:': [{'Before Wash': ['18.2', '18.8'], 'After Wash': ['17.8', '15.5'],
                           'Dimensional Change': ['-0.55', '0.22']}],
            'Sample C:': [{'Before Wash': ['18.2', '18.8'], 'After Wash': ['17.8', '15.5'],
                           'Dimensional Change': ['-0.55', '0.22']}],
            'Sample D:': [{'Before Wash': ['18.2', '18.8'], 'After Wash': ['17.8', '15.5'],
                           'Dimensional Change': ['-0.55', '0.22']}]}
    return test_name,test_method,remarks,data

