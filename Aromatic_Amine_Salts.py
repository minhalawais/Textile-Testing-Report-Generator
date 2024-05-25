from create_test import create_test
from docx.enum.section import WD_SECTION
from docx.shared import Inches,Pt
from global_func import set_col_widths

def Aromatic_test_template(doc,test_name,test_method,remarks,values,req):
    create_test(doc,test_name,test_method,remarks)
    # Define the maximum number of dictionary columns per table
    max_columns_per_table = 6

    num_tables = (len(values) + max_columns_per_table - 1) // max_columns_per_table



    for table_index in range(num_tables):
        start_column = table_index * max_columns_per_table
        end_column = start_column + max_columns_per_table

        num_columns = 1 + min(max_columns_per_table, len(values) - start_column)
        table = doc.add_table(rows=2, cols=num_columns+1)
        table.style = 'Table Grid'
        table.cell(1, 0).text = 'Result'

        for column_index, column_title in enumerate(list(values.keys())[start_column:end_column], start=1):
            table.cell(0, column_index).text = column_title

            row = table.rows[0]  # Start from the second row since the first row contains the headers
            row1 = table.rows[1]
            for column_index, column_values in enumerate(list(values.values())[start_column:end_column], start=1):
                cell = row.cells[column_index]
                cell.text = list(values.keys())[start_column+(column_index-1)]
                cell = row1.cells[column_index]
                cell.text = column_values
        a = table.cell(0,num_columns)
        a.text = "mg/kg"
        b = table.cell(1,num_columns)
        b.text = req[0]
        widths = (Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5),Inches(1.5), Inches(1.5), Inches(1.5))
        set_col_widths(table,widths)
        sample_heading = doc.add_paragraph()
        sample_heading.paragraph_format.space_before = Pt(0)
        sample_heading.paragraph_format.space_after = Pt(0)
    table = doc.add_table(rows=5, cols=1)
    table.style = 'Table Grid'
    table.cell(0,0).text ="Note"
    table.cell(1,0).text = "n.d. = not detected"
    table.cell(2,0).text = "mg/kg = ppm"
    table.cell(3,0).text = "* = Exceeds the limit	"
    table.cell(4,0).text = "Detection Limit = 5 mg/kg (for individual compound)	"
    widths = (Inches(10), Inches(2))
    set_col_widths(table, widths)
    data = [
        (1, "4-Chloro-o-toluidinium chloride", "3165-93-3", 3, "4-meth methoxy-m-phenylene diammonium  sulphate; \n2,4- diaminoanisole sulphate", "120-71-8"),
        (2, "2-Naphthylammoniumacetate", "553-00-4", 4, "2,4,5-trimethylaniline hydrochloride", "21436-97-5")]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    runner = p.add_run("List of Aromatic Amine Salts:")
    runner.bold = True
    runner.italic = True

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Set table header and bold the first row
    headings = ['Sr#', 'Substance name', 'CAS No.', 'Sr#', 'Substance name', 'CAS No.']
    for i, heading in enumerate(headings):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        paragraph.text = heading
        run = paragraph.runs[0]
        run.font.bold = True

    # Insert data into the table
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = str(cell_value)

    # Adjust column widths
    for col in table.columns:
        for cell in col.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size (optional)
    widths = (Inches(0.3), Inches(2.5), Inches(1.2), Inches(0.3), Inches(2.5), Inches(1.2))
    set_col_widths(table, widths)

def Aromatic_test_data():

    values = {
        'A1+A2+A3': 'ND',
        'B1+B2+B3': 'ND',
        'C1+C2+C3': 'ND',
        'D1+D2+D3': 'ND',
        'E1+E4+E5': 'ND',
        'E2+E3': 'ND',
        'F1+F2+F3': 'ND',
        'F4+F5': 'ND',
        'G1+G3+G5': 'ND',
        'G2+G4': 'ND',
        'H1+H2+H3': 'ND'
    }
    test_name = "Aromatic Amine Salts:"
    test_method = "All Textile: According to DIN EN ISO 14362-1:2017 – Analysis was conducted with GC-MS/HPLC-DAD. "
    remarks = ""
    return test_name,test_method,remarks,values

