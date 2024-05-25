from create_test import create_test
from docx.enum.section import WD_SECTION
from docx.shared import Inches,Pt
from global_func import set_col_widths

def Ph_test_template(doc,test_name,test_method,remarks,values,req):
    create_test(doc,test_name,test_method,remarks)
    # Define the maximum number of dictionary columns per table
    max_columns_per_table = 6

    num_tables = (len(values) + max_columns_per_table - 1) // max_columns_per_table



    for table_index in range(num_tables):
        start_column = table_index * max_columns_per_table
        end_column = start_column + max_columns_per_table

        num_columns = 1 + min(max_columns_per_table, len(values) - start_column)
        table = doc.add_table(rows=2, cols=num_columns+1)
        table.style = 'Table Grid'
        table.cell(1, 0).text = 'Result'
        for column_index, column_title in enumerate(list(values.keys())[start_column:end_column], start=1):
            table.cell(0, column_index).text = column_title

            row = table.rows[0]  # Start from the second row since the first row contains the headers
            row1 = table.rows[1]
            for column_index, column_values in enumerate(list(values.values())[start_column:end_column], start=1):
                cell = row.cells[column_index]
                cell.text = list(values.keys())[start_column+(column_index-2)]
                cell = row1.cells[column_index]
                cell.text = column_values
        b = table.cell(0, num_columns)
        b.text = "mg/kg"
        b = table.cell(1,num_columns)
        b.text = req[0]
        widths = (Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5),Inches(1.5), Inches(1.5), Inches(1.5))
        set_col_widths(table,widths)
        sample_heading = doc.add_paragraph()
        sample_heading.paragraph_format.space_before = Pt(0)
        sample_heading.paragraph_format.space_after = Pt(0)
    table = doc.add_table(rows=5, cols=1)
    table.style = 'Table Grid'
    table.cell(0,0).text =" Note"
    table.cell(1,0).text = " N. D. = Not detected"
    table.cell(2, 0).text = " Laboratory Reporting Limit:  50 mg/kg"
    table.cell(3, 0).text = " N. D. = Not detected"
    table.cell(4, 0).text = " mg/kg: milligram per kilogram"
    widths = (Inches(10), Inches(2))
    set_col_widths(table, widths)
    data = [
        (1, "Bis (2-ethylhexyl) phthalate (DEHP)", "117-81-7", 12,
         "1,2-Benzendicarboxylicacidalkyl esters, C7-rich (DIHP)", "71888-89-6"),
        (2, "Dibutyl phthalate (DBP)", "84-74-2", 13,
         "1,2-Benzenedicarboxylic acid, dihexyl ester, branched and linear (DHxP)", "68515-50-4"),
        (3, "Benzyl butyl phthalate (BBP)", "85-68-7", 14, "Dimethyl phthalate (DMP)", "131-11-3"),
        (4, "Diisobutyl phthalate (DIBP)", "84-69-5", 15, "Di-n-propyl phthalate (DPP)", "131-16-8"),
        (5, "Di-“isononyl” phthalate (DINP)", "28553-12-0\n68515-48-0", 16, "Dicyclohexyl phthalate (DCP)",
         "84-61-7\n55819-02-8\n169741-16-6"),
        (6, "Di-“isodecyl” phthalate (DIDP)", "26761-40-0\n68515-49-1", 17,
         "1,2-Benzenedicarboxylic acid, di-2-propenyl ester (DAP)", "131-17-9"),
        (7, "Di-n-octyl phthalate (DNOP)", "117-84-0", 18, "Di-iso-hexylphthalate (DIHxP)", "71850-09-4"),
        (8, "Di-n-pentylphthalate (n-, iso-, or mixed) (DIPP/ DNPP)", "131-18-0\n605-50-5\n776297-69-9\n84777-06-0", 19,
         "1,2-Benzenedicarboxylic acid, di-C6-10 alkyl esters", "68515-51-5"),
        (9, "Bis (2-methoxyethyl) phthalate (DMEP)", "117-82-8", 20,
         "1,2-Benzenedicarboxylic acid, mixed decyl and hexyl and octyl diesters", "68648-93-1"),
        (10, "Di-n-hexyl phthalate (DNHP)", "84-75-3", 21, "Di-ethylphthalate (DEP)", "84-66-2"),
        (11.1, "1,2-Benzendicarboxylicacid, di-C7-11branched and linear alkylesters (DHNUP)", "68515-42-4", 22,
         "1,2-Cyclohexane dicarboxylic acid diisononyl ester (DINCH)", "166412-78-8"),
        (11.2, "Di-2-propyl heptyl phthalate (DPHP)", "53306-54-0", 23, "", ""),
        (11.3, "Di-n-nonylphthalate (DNP)", "84-76-4", 24, "", ""),
        (11.4, "Diisooctyl phthalate (DIOP)", "27554-26-3", 25, "", "")
    ]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    runner = p.add_run("List of Phthalates:")
    runner.bold = True
    runner.italic = True

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Set table header and bold the first row
    headings = ['Sr#', 'Substance name', 'CAS No.', 'Sr#', 'Substance name', 'CAS No.']
    for i, heading in enumerate(headings):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        paragraph.text = heading
        run = paragraph.runs[0]
        run.font.bold = True

    # Insert data into the table
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = str(cell_value)

    # Adjust column widths
    for col in table.columns:
        for cell in col.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Adjust font size (optional)
    widths = (Inches(0.3), Inches(2.5), Inches(1.2), Inches(0.3), Inches(2.5), Inches(1.2))
    set_col_widths(table, widths)


def Ph_test_data():

    values = {
        'A1+A2+A3': 'ND',
        'B1+B2+B3': 'ND',
        'C1+C2+C3': 'ND',
        'D1+D2+D3': 'ND',
        'E1+E4+E5': 'ND',
        'E2+E3': 'ND',
        'F1+F2+F3': 'ND',
        'F4+F5': 'ND',
        'G1+G3+G5': 'ND',
        'G2+G4': 'ND',
        'H1+H2+H3': 'ND'
    }
    test_name = "Phthalates"
    test_method = "DIN 54231:2005, Analyzed by GC-MS"
    remarks = ""
    return test_name,test_method,remarks,values

