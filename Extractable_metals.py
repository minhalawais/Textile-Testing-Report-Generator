from create_test import create_test
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from global_func import set_col_widths

def extract_test_template(doc,test_name,test_method,remarks,title,cas_no,mdl,values,req):
    create_test(doc,test_name,test_method,remarks)
    # Define the maximum number of dictionary columns per table
    max_columns_per_table = 5
    last = ['0.2','0.2','0.1','0.5','30','1.0','1.0','25','1.0','0.02','1000','100']
    num_tables = (len(values) + max_columns_per_table) // max_columns_per_table



    for table_index in range(num_tables):
        if table_index % 2 == 0 and table_index!=0:
            doc.add_section(WD_SECTION.NEW_PAGE)
            create_test(doc,test_name,test_method,remarks)

        start_column = table_index * max_columns_per_table
        end_column = start_column + max_columns_per_table

        num_columns = 2 + min(max_columns_per_table, len(values) - start_column)
        table = doc.add_table(rows=len(title)+1, cols=num_columns+2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Substance Name'
        table.cell(0, 1).text = 'CAS No.'
        table.cell(0, 2).text = 'MDL (mg/kg)'
        table.cell(0,num_columns+1).text = "(mg/kg)"

        for column_index, column_title in enumerate(list(values.keys())[start_column:end_column], start=3):
            table.cell(0, column_index).text = column_title

        for i, (value, cas,mdl_value,last_value) in enumerate(zip(title, cas_no,mdl,req)):

            row = table.rows[i + 1]  # Start from the second row since the first row contains the headers

            cell1 = row.cells[0]
            cell1.text = value

            cell2 = row.cells[1]
            cell2.text = cas

            cell3 = row.cells[2]
            cell3.text = mdl_value
            cell3 = row.cells[num_columns+1]
            cell3.text = last_value

            for column_index, column_values in enumerate(list(values.values())[start_column:end_column], start=3):
                cell = row.cells[column_index]
                cell.text = column_values[i]

        sample_heading = doc.add_paragraph()
        widths = (Inches(2),  Inches(1), Inches(0.5), Inches(0.5),Inches(0.5),Inches(0.5),Inches(0.5),Inches(0.5),Inches(0.5),Inches(3))
        set_col_widths(table, widths)
    table = doc.add_table(rows=4, cols=1)
    table.style = 'Table Grid'
    table.cell(0,0).text ="Note"
    table.cell(1,0).text = "Method Detection Limit"
    table.cell(2,0).text = "mg/kg: milligram per kilogram"
    table.cell(3,0).text = "ND: Not detected"
    widths = (Inches(10), Inches(1))
    set_col_widths(table, widths)
def extract_test_data():
    title = ['As (Arsenic)','Pb (Lead)','Cd (Cadmium)','Cr VI (Chromium VI)','Sb (Antimony)','Cr (Chromium)','Co (Cobalt)','Cu (Copper)','Ni (Nickel)','Hg (Mercury)','Ba (Barium)','Se (Selenium)']

    cas_no = ['7440-38-2'  ,'7439-92-1','7440-43-9','18540-29-9','7440-36-0','7440-47-3','7440-48-4','7440-50-8','7440-02-0','7439-97-6','7440-39-3','7782-49-2']
    mdl = ['0.2',
'0.2',
'0.1',
'0.5',
'0.5',
'0.5',
'0.5',
'0.5',
'0.5',
'0.02',
'0.5',
'0.5'
]
    values = {
        'Sample A': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample B': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample C': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample D': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample E': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample F': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample G': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample H': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample I': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample J': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND'],
        'Sample K': ['ND', 'ND', 'ND', 'ND', 'ND','ND', 'ND', 'ND', 'ND', 'ND','ND','ND']
    }
    test_name = "Extractable (heavy) Metals:"
    test_method = "DIN EN 16711-2:2016, Analyzed by ICP-OES and DIN EN ISO 17075-1:2017 (modified) for Cr (VI)"
    remarks = ""
    return test_name,test_method,remarks,title,cas_no,mdl,values