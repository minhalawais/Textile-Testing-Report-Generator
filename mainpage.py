from docx import Document
from docx.shared import Pt,RGBColor,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from generate_qrcode import generate_qr_code_with_image

from global_func import requirement_data

doc = Document()

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size=Pt(10)
font_name = 'Calibri'
font_size = Pt(10)





def set_col_widths(table,widths):

    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

#first table
table = doc.add_table(rows=5, cols=5)
table.style = 'Table Grid'



applicant_list = ['Applicant:','Contact:','Address:','Tel:','E-mail:']
applicant_dict = {'Applicant:':'Kamal Mills Private Limited.','Contact:':'Muhammad Asad','Address:':'3 KM, Jhumra Road, Khurrianwala, Faisalabad-Pakistan.','Tel:':'0300-8793742','E-mail:':'Muhammad.asad@kamal.com.pk'}
buyer_list = ['Buyer Name:','Supplier Name:','Agent:','Country of Origin:','Country of Destination:']
buyer_dict = {'Buyer Name:':'Kik Textilien und Non-Food GmbH','Supplier Name:':'Kamal Mills Private Limited.','Agent:':'Matrix Sourcing','Country of Origin:':'Pakistan','Country of Destination:':'Germany'}
sample_list = ['Product Description:','Material Composition:','Fabric:','Fabric Weight:','Merchandise Category (WGR):','P.O Reference No:','P.O:','Article No:','Article Description:','Style Color:','Supplier No:','Season:','Buying Dept (EKB) :','Dye stuff:','Previous Report# (for retest):']
sample_dict = {'Product Description:':'Children boy ergee socks 4 pair','Material Composition:':'  72% Cotton, 26% Polyamide, 2% Elastane','Fabric:':'Jersey  ','Fabric Weight:':'/','Merchandise Category (WGR):':'837','P.O Reference No:':'P203649','P.O:':'4500325875,4500325876','Article No:':'1165292903,1165292904','Article Description:':'Children boy ergee socks 4 pair','Style Color:':'Orange, Green','Supplier No:':'301313','Season:':'123','Buying Dept (EKB) :':'/','Dye stuff:':'/','Previous Report# (for retest):':'/'}
physical_test = {}
chemical_test = {}
for i in range(5):
    row = table.rows[i]
    row_cells = row.cells[0].paragraphs[0].add_run(applicant_list[i])
    row_cells.bold = True
    row_cells.font.size = Pt(10)
    row_cells1 = row.cells[1].paragraphs[0].add_run(applicant_dict[applicant_list[i]])
    row_cells = row.cells[3].paragraphs[0].add_run(buyer_list[i])
    row_cells.bold = True
    row_cells.font.size = Pt(10)
    row_cells = row.cells[4].paragraphs[0].add_run(buyer_dict[buyer_list[i]])


sample_heading = doc.add_paragraph()

widths = (Inches(1), Inches(3), Inches(0.2),Inches(1.5), Inches(3))
set_col_widths(table,widths)

#second table
table = doc.add_table(rows=16, cols=3)
table.style = 'Table Grid'

for i, item in enumerate(sample_list,start=1):
    cell1 = table.cell(i, 1)
    cell2 = table.cell(i, 2)

    row_cells = cell1.paragraphs[0].add_run(item)
    row_cells.bold = True
    row_cells.font.size = Pt(10)
    row_cells1 = cell2.paragraphs[0].add_run(sample_dict.get(item, ''))
a = table.cell(1,0)
for j in range(1, 16):
    b = table.cell(j, 0)
    A = a.merge(b)
    a = A
a = table.cell(0,0)
b = table.cell(0, 1)
c = table.cell(0, 2)
A = a.merge(b)
B=A.merge(c)
shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w')))
B._tc.get_or_add_tcPr().append(shading_elm_1)
B.text = "Sample Information	"
cell = table.cell(1, 0)
paragraph = cell.paragraphs[0]

paragraph = cell.add_paragraph()
run = paragraph.add_run()
run.add_picture(r"Logo\sample.png", width=Inches(2))
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sample_heading = doc.add_paragraph()
widths = (Inches(1.5), Inches(2), Inches(4))
set_col_widths(table,widths)

#third table
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
widths = (Inches(3), Inches(8))
set_col_widths(table,widths)
cell1 = table.cell(0, 0)
cell2 = table.cell(1, 0)
row_cells = cell1.paragraphs[0].add_run("Submitted Care instructions:")
row_cells.bold = True
row_cells.font.size = Pt(10)
row_cells = cell2.paragraphs[0].add_run("Test Package:	")
row_cells.bold = True
row_cells.font.size = Pt(10)
cell1 = table.cell(0, 1)
cell2 = table.cell(1, 1)
cell1.text = "/"
cell2.text = '“PC A: (Socks size > 24)”	'
#Generate QrCode
text = f"Report# 02249-23/ Issued Date: 06 02 02/ Rating: Pass/ Applicant: {applicant_dict['Applicant:']}/ Client: {buyer_dict['Buyer Name:']}/ Sample: {sample_dict['Product Description:']}"
image_path = r"Logo\tti_logo.png"

generate_qr_code_with_image(text, image_path, )

#fourth table
sample_heading = doc.add_paragraph()
sample_run = sample_heading.add_run("\n\n")
table = doc.add_table(rows=1, cols=3)
cell1 = table.cell(0,0)
paragraph = cell1.paragraphs[0]
paragraph.alignment = 0
run = paragraph.add_run()
run.add_picture(r"Logo\qrcode.png", width=Inches(1.5))
paragraph.alignment = 0
cell3 = table.cell(0,2)
paragraph = cell3.paragraphs[0]
sample_run = paragraph.add_run("For and on behalf of\nTEXTILE TESTING INTERNATIONAL")
paragraph = cell3.add_paragraph()
paragraph.alignment = 1
sample_run = paragraph.add_run()
sample_run.add_picture(r"Logo\signature.png", width=Inches(2))
paragraph = cell3.add_paragraph()
sample_run = paragraph.add_run("Ali Ashraf\n    AVP Softlines ")
sample_run.bold = True
paragraph.alignment = 1
widths = (Inches(2.59), Inches(3), Inches(2.18))
set_col_widths(table, widths)
doc.add_section(WD_SECTION.NEW_PAGE)
for row in table.rows:
    row.height = Inches(1.4)
#test names page
sample_run.font.size = Pt(8)
'''
from test_name_page import test_name_table
test_name_table(doc)
doc.add_section(WD_SECTION.NEW_PAGE)
sample_heading = doc.add_paragraph()
sample_run = sample_heading.add_run("\n")
'''
#component breakdown page
from component_breakdown import breakdown_table
breakdown_table(doc)
doc.add_section(WD_SECTION.NEW_PAGE)
sample_heading = doc.add_paragraph()
sample_run = sample_heading.add_run("\n")

#collage image

from images_collage import set_image
set_image(doc)
doc.add_section(WD_SECTION.NEW_PAGE)
sample_heading = doc.add_paragraph()
sample_run = sample_heading.add_run("\n")



#dimensional change
from Dimensional_stability_to_washing import dimensional_test_data,dimensional_test_template
test_name,test_method,remarks,data = dimensional_test_data()
images = ['../static/images/test_template.jpg', '../static/images/current_page.jpg', '../static/images/page_1.jpg']
input_data_list = [["±5", "±7", "±8", "±10"]]
input_data_list = requirement_data(input_data_list)
dimensional_test_template(doc, test_name, test_method, remarks, data, '+-5')
doc.add_section(WD_SECTION.NEW_PAGE)
sample_heading = doc.add_paragraph()


#appearance after washing
from Appearance_after_washing import appearance_test_data,appearance_test_template
test_name,test_method,remarks,title,req,sample_results = appearance_test_data()
appearance_test_template(doc,test_name,test_method,remarks,title,req,sample_results)
doc.add_section(WD_SECTION.NEW_PAGE)
sample_heading = doc.add_paragraph()




#colour fastness to water
from Colour_fastness_to_water import color_test_data,color_test_template
test_name,test_method,remarks,title,values = color_test_data()
color_test_template(doc,test_name,test_method,remarks,title,values)
doc.add_section(WD_SECTION.NEW_PAGE)
sample_heading = doc.add_paragraph()



#colour fastness to perspiration
from Colour_fastness_to_perspiration import color_fp_test_data,color_fp_test_template
test_name,test_method,remarks,title,values = color_fp_test_data()
color_fp_test_template(doc,test_name,test_method,remarks,title,values)
doc.add_section(WD_SECTION.NEW_PAGE)
sample_heading = doc.add_paragraph()




#colour fastness to water
from Colour_fastness_to_rubbing import Color_fr_test_data,Color_fr_test_template
test_name,test_method,remarks,title,values = Color_fr_test_data()
Color_fr_test_template(doc,test_name,test_method,remarks,title,values)
doc.add_section(WD_SECTION.NEW_PAGE)
sample_heading = doc.add_paragraph()
sample_run = sample_heading.add_run("\n")


#fiber composition
from fiber_composition import fiber_test_template,fiber_test_data
test_name = "Fiber Composition:"
test_method = "DIN EN ISO 1833-7:2017,11:2017"
remarks = ''
test_name,test_method,remarks,title,sample_results = fiber_test_data()
fiber_test_template(doc,test_name,test_method,remarks,title,sample_results)
doc.add_section(WD_SECTION.NEW_PAGE)



#Alkylphenole test
from Alkylphenole import Alkylphenole_test_data,Alkylphenole_test_template
test_name,test_method,remarks,title,cas_no,values = Alkylphenole_test_data()
Alkylphenole_test_template(doc,test_name,test_method,remarks,title,cas_no,values)
doc.add_section(WD_SECTION.NEW_PAGE)


#Fabric Weight
from fabric_weight import fabric_test_data,fabric_test_template
test_name,test_method,remarks,title,values = fabric_test_data()
fabric_test_template(doc,test_name,test_method,remarks,title,values)
doc.add_section(WD_SECTION.NEW_PAGE)


#Seam spirality after laundering:
from Seam_spirality_after_laundering import seam_test_template,seam_test_data
test_name,test_method,remarks,title,values = seam_test_data()
seam_test_template(doc,test_name,test_method,remarks,title,values)
doc.add_section(WD_SECTION.NEW_PAGE)

#Azo-Dyes (including Aniline):
from Azo_Dyes import Azo_test_template,Azo_test_data
test_name,test_method,remarks,values = Azo_test_data()
Azo_test_template(doc,test_name,test_method,remarks,values)
doc.add_section(WD_SECTION.NEW_PAGE)

#Aromatic Amine Salts:
from Aromatic_Amine_Salts import Aromatic_test_data,Aromatic_test_template
test_name,test_method,remarks,values = Aromatic_test_data()
Aromatic_test_template(doc,test_name,test_method,remarks,values)
doc.add_section(WD_SECTION.NEW_PAGE)

#Formaldehyde
from Formaldehyde import Formaldehyde_test_data,Formaldehyde_test_template
test_name,test_method,remarks,values = Formaldehyde_test_data()
Formaldehyde_test_template(doc,test_name,test_method,remarks,values)
doc.add_section(WD_SECTION.NEW_PAGE)

#Total Lead
from Total_Lead_Content import Total_test_data,Total_test_template
test_name,test_method,remarks,values = Total_test_data()
Total_test_template(doc,test_name,test_method,remarks,values)
doc.add_section(WD_SECTION.NEW_PAGE)
#Total Cadmium (Cd) Content:
from Total_Cadmium import Cadmium_test_template,Cadmium_test_data
test_name,test_method,remarks,values = Cadmium_test_data()
Cadmium_test_template(doc,test_name,test_method,remarks,values)
doc.add_section(WD_SECTION.NEW_PAGE)
#Extractable Metals
from Extractable_metals import extract_test_data,extract_test_template
test_name,test_method,remarks,title,cas_no,mdl,values = extract_test_data()
extract_test_template(doc,test_name,test_method,remarks,title,cas_no,mdl,values)
doc.add_section(WD_SECTION.NEW_PAGE)

#Quinoline
from Quinoline import Quinoline_test_template,Quinoline_test_data
test_name,test_method,remarks,values =Quinoline_test_data()
Quinoline_test_template(doc,test_name,test_method,remarks,values)
doc.add_section(WD_SECTION.NEW_PAGE)


#Phthalates
from phthalates import Ph_test_data,Ph_test_template
test_name,test_method,remarks,values = Ph_test_data()
Ph_test_template(doc,test_name,test_method,remarks,values)


#Polycyclic
from Polycyclic import Poly_test_data,Poly_test_template
test_name,test_method,remarks,title,cas_no,values = Poly_test_data()
Poly_test_template(doc,test_name,test_method,remarks,title,cas_no,values)
from header import set_header
set_header(doc)
doc.save("test_template.docx")



#set footer
from footer import  set_footer
set_footer()

