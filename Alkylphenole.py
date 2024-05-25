from create_test import create_test
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from global_func import set_col_widths

def Alkylphenole_test_template(doc,test_name,test_method,remarks,title,cas_no,values,req):
    create_test(doc,test_name,test_method,remarks)
    # Define the maximum number of dictionary columns per table
    max_columns_per_table = 4

    num_tables = (len(values) + max_columns_per_table - 1) // max_columns_per_table



    for table_index in range(num_tables):


        start_column = table_index * max_columns_per_table
        end_column = start_column + max_columns_per_table

        num_columns = 2 + min(max_columns_per_table, len(values) - start_column)
        table = doc.add_table(rows=len(title)+1, cols=num_columns+1)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Title'
        table.cell(0, 1).text = 'CAS No.'

        for column_index, column_title in enumerate(list(values.keys())[start_column:end_column], start=2):
            table.cell(0, column_index).text = column_title

        for i, (value, cas) in enumerate(zip(title, cas_no)):

            row = table.rows[i + 1]  # Start from the second row since the first row contains the headers

            cell1 = row.cells[0]
            cell1.text = value

            cell2 = row.cells[1]
            cell2.text = cas

            for column_index, column_values in enumerate(list(values.values())[start_column:end_column], start=2):
                cell = row.cells[column_index]
                cell.text = column_values[i]
        a = table.cell(1,num_columns).merge(table.cell(2,num_columns))
        a.text = req[0]
        b = table.cell(3, num_columns).merge(table.cell(4, num_columns))
        b.text = req[1]
        table.cell(5,num_columns).text = req[2]
        if (table_index +1)%2 == 0 and table_index!=0:
            doc.add_section(WD_SECTION.NEW_PAGE)
            create_test(doc,test_name,test_method,remarks)
        sample_heading = doc.add_paragraph()
        widths = (Inches(1.5), Inches(1.5), Inches(1), Inches(1), Inches(1), Inches(1))
        set_col_widths(table, widths)
        for row in table.rows:
            row.cells[-1].width = Inches(2)
    table = doc.add_table(rows=4, cols=1)
    table.style = 'Table Grid'
    table.cell(0,0).text ="Note"
    table.cell(1,0).text = "Unit: mg/kg: (milligram per kilogram)	"
    table.cell(2,0).text = "MDL: 5mg/kg for (NP/OP): 30mg/kg (NPEO/OPEO)"
    table.cell(3,0).text = "ND: Not detected	"
    widths = (Inches(10), Inches(1))
    set_col_widths(table, widths)
def Alkylphenole_test_data():
    title = [
        'Nonylphenol (NP), mixed isomers',
        'Octylphenol (OP), mixed isomers',
        'Nonylphenol Ethoxylates (NPEO)',
        'Octylphenol Ethoxylates (OPEO)',
        'Sum NP/NPEO'
    ]

    cas_no = [
        "104-40-5\n11066-49-2\n25154-52-3\n84852-15-3",
        "140-66-9\n1806-26-4\n27193-28-8",
        "9016-45-9\n26027-38-3\n37205-87-1\n68412-54-4\n127087-87-0",
        "9002-93-1\n9036-19-5\n68987-90-6",
        "-"
    ]

    values = {
        'Sample A': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample B': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample C': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample D': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample E': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample F': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample G': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample H': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample I': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample J': ['ND', 'ND', 'ND', 'ND', 'ND'],
        'Sample K': ['ND', 'ND', 'ND', 'ND', 'ND']
    }
    test_name = "Alkylphenole/ Alkylphenolethoxylate (AP/APEO):"
    test_method = "EN ISO 21084:2019/ AP: Analyzed by GC-MS / EN ISO 18254-1:2016 / APEO: \nAnalyzed by LC-MS"
    remarks = ""
    return test_name,test_method,remarks,title,cas_no,values