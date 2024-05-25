from docx import Document
from docx.shared import Pt,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from global_func import set_row_color,set_col_widths
def copy_table(source_table, dest_doc):
    new_table = dest_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    new_table.style = 'Table Grid'
    set_row_color(new_table.rows[0])
    widths = (Inches(2.5), Inches(2.5), Inches(4), Inches(3))
    set_col_widths(new_table,widths)
    # Create a list to keep track of the text in the first column
    first_column_text = []

    for row_idx, source_row in enumerate(source_table.rows):
        for col_idx, source_cell in enumerate(source_row.cells):
            new_cell = new_table.cell(row_idx, col_idx)
            new_cell.width = source_cell.width
            new_cell.paragraphs[0].clear()  # Clear any existing content
            new_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # If it's the first column, process merging
            if col_idx == 0:
                # Get the text from the first cell in the source row
                cell_text = source_cell.text

                if row_idx > 0 and cell_text == first_column_text[-1]:
                    # If the current cell text matches the previous cell text, merge the cells
                    first_cell = new_table.cell(row_idx, 0)
                    previous_cell = new_table.cell(row_idx - 1, 0)
                    merged_cell = first_cell.merge(previous_cell)
                    merged_cell.text = cell_text
                    # Center align the merged cell's text
                    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # If the current cell text is different, add it to the list
                    first_column_text.append(cell_text)

                # Add the text only once to the merged cell
                if col_idx == 0:
                    new_cell.paragraphs[0].add_run(cell_text).font.size = Pt(10)
                else:
                    # Ensure that non-merged cells are also center-aligned
                    new_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            else:
                # If it's not the first column, simply copy the cell text and formatting
                new_cell.paragraphs[0].add_run(source_cell.text).font.size = Pt(10)
                # Center align non-merged cells
                new_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Copy cell background color
            if source_cell._tc is not None:
                cell_bg_color = source_cell._tc.find('.//w:shd', namespaces=source_cell._tc.nsmap)
                if cell_bg_color is not None:
                    new_cell._tc.append(cell_bg_color)

            # Copy merged cell properties
            if source_cell._tc is not None:
                cell_vMerge = source_cell._tc.find('.//w:vMerge', namespaces=source_cell._tc.nsmap)
                if cell_vMerge is not None:
                    new_cell._tc.append(cell_vMerge)

def breakdown_table(dest_doc):
    # Example usage:
    source_docx_path = "component_format.docx"
    table_index = 1  # Index of the table to copy (0-based)

    source_doc = Document(source_docx_path)
    if table_index < len(source_doc.tables):
        source_table = source_doc.tables[table_index]
        copy_table(source_table, dest_doc)

        print(f"Table {table_index} copied and pasted to the destination document.")
    else:
        print(f"Table {table_index} not found in the source document.")

