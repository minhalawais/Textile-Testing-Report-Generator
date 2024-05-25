from docx import Document
from docx.shared import Pt,RGBColor,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from global_func import set_row_color,set_col_widths
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w')))

'''
def test_name_table(doc,test_name,result):
    physical_test = ['Dimensional stability to washing', 'Appearance after washing', 'Colour fastness to water', 'Colour fastness to perspiration', 'Colour fastness to rubbing', 'Fiber Composition']
    chemical_test = ['Alkylphenole/ Alkylphenolethoxylate (AP/APEO)','Azo-Dyes (including Aniline)','Aromatic Amine Salts']
    #pass = {'Physical Tests':"Pass", 'Dimensional stability to washing':'Pass', 'Appearance after washing':"Pass", 'Colour fastness to water', 'Colour fastness to perspiration', 'Colour fastness to rubbing', 'Fiber Composition'}
    header = ['Sr. No','Test properties','PASS','FAIL','Remarks']
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    set_row_color(table.rows[0])
    for i in range(5):
        row_cells = hdr_cells[i].paragraphs[0].add_run(header[i])
        hdr_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER


        row_cells.bold = True
        row_cells.font.size = Pt(10)
    row_cells = table.add_row().cells
    a = table.cell(1, 0)
    for i in range(5):
        b = table.cell(1,i)
        A = a.merge(b)
        a = A
    row_cells = a.paragraphs[0].add_run("Physical Tests")
    row_cells.bold = True
    row_cells.font.size = Pt(10)
    for i in range(1, len(physical_test)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[1].text = physical_test[i]
        row_cells[2].text = ''
        row_cells[3].text = ''
        row_cells[4].text = ''
    row_cells = table.add_row().cells
    a = row_cells[0]
    for i in range(5):
        b = row_cells[i]
        A = a.merge(b)
        a = A
    row_cells = a.paragraphs[0].add_run("Chemical Tests")
    row_cells.bold = True
    row_cells.font.size = Pt(10)
    for i in range(0, abs(len(chemical_test)-len(physical_test))):
        row_cells = table.add_row().cells
        row_cells[0].text = str(abs(i+len(physical_test)))
        row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[1].text = chemical_test[i]
        row_cells[2].text = ''
        row_cells[3].text = ''
        row_cells[4].text = ''
    widths = (Inches(0.9), Inches(4), Inches(0.5),Inches(0.5), Inches(4))
    set_col_widths(table,widths)
    table.rows[1].cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
    table.rows[len(physical_test)+1].cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        row.height = Inches(0.25)'''
def test_name_table(doc):
    physical_test = ['Dimensional stability to washing', 'Appearance after washing', 'Colour fastness to water', 'Colour fastness to perspiration', 'Colour fastness to rubbing', 'Fiber Composition']
    chemical_test = ['Alkylphenole/ Alkylphenolethoxylate (AP/APEO)','Azo-Dyes (including Aniline)','Aromatic Amine Salts']
    #pass = {'Physical Tests':"Pass", 'Dimensional stability to washing':'Pass', 'Appearance after washing':"Pass", 'Colour fastness to water', 'Colour fastness to perspiration', 'Colour fastness to rubbing', 'Fiber Composition'}
    header = ['Sr. No','Test properties','PASS','FAIL','Remarks']
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    set_row_color(table.rows[0])
    for i in range(5):
        row_cells = hdr_cells[i].paragraphs[0].add_run(header[i])
        hdr_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER


        row_cells.bold = True
        row_cells.font.size = Pt(10)
    widths = (Inches(0.9), Inches(4), Inches(0.5), Inches(0.5), Inches(4))
    set_col_widths(table, widths)