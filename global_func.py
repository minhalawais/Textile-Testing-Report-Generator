from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx import Document
from docx.shared import Pt,RGBColor,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx import oxml
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
from PIL import Image
from PIL import ImageDraw
from PIL import ImageFont
import pythoncom
import docx2pdf
from generate_qrcode import generate_qr_code_with_image

def convert_to_pdf():
    try:
        #pythoncom.CoInitialize()  # Initialize the COM library
        output_path = 'static/test_template.pdf'
        docx2pdf.convert('test_template.docx', output_path)

        print(f"Conversion successful. PDF saved at {output_path}")
    except Exception as e:
        print(f"An error occurred: {e}")
def set_row_color(row):
    for i in row.cells:
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w')))
        i._tc.get_or_add_tcPr().append(shading_elm_1)
def set_col_widths(table,widths):

    for row in table.rows:
        for idx, width in enumerate(widths):
            try:
                row.cells[idx].width = width
            except:
                pass
def requirement_data(options):
    input_data_list = []
    for i in range(len(options)):
        input_data_list.append(
            {
                'id': f'input{i+1}',
                'label': f'Input {i+1}',
                'datalist_id': f'datalist{i+1}',
                'options': options[i]
            }
        )
    return input_data_list
import os
from docx2pdf import convert
from pdf2image import convert_from_path



def insert_image_in_docx_table(doc):
    docx_file = 'test_template.docx'
    image_path = 'static/edited_front_image.png'
    # Load the docx file

    # Access the second table (assuming it exists)
    if len(doc.tables) > 1:
        table = doc.tables[1]

        # Access the cell in the second row of the first column
        cell = table.cell(1, 0)

        # Insert an image into the cell
        img = Image.open(image_path)
        width, height = img.size
        aspect_ratio = height / width
        new_width = Inches(2)  # You can adjust this width according to your preference
        new_height = Inches(2 * aspect_ratio)
        cell.paragraphs[0].add_run().add_picture(image_path, width=new_width, height=new_height)

        # Save the modified document
        doc.save("test_template.docx")
        print("Image added successfully")
    else:
        print("Table 2 does not exist in the document.")

def add_image_to_docx(doc):
    docx_file = 'test_template.docx'
    image_path = 'static/edited_image.png'
    # Load the docx file

    # Access the second table (assuming it exists)
    if len(doc.tables) > 1:
        table = doc.tables[6]

        # Access the cell in the second row of the first column
        cell = table.cell(0, 0)

        # Insert an image into the cell
        img = Image.open(image_path)
        width, height = img.size
        aspect_ratio = height / width
        new_width = Inches(2)  # You can adjust this width according to your preference
        new_height = Inches(2 * aspect_ratio)
        cell.paragraphs[0].add_run().add_picture(image_path,width=Inches(6.5),height=Inches(8))

        # Save the modified document
        doc.save("test_template.docx")
        print("Image added successfully")
    else:
        print("Table 2 does not exist in the document.")
def doc_to_jpg(specific_pages):
    word_file = r'test_template.docx'
    output_folder = r'static\images'

    pdf_file = "temp.pdf"
    # Set output image format
    convert(word_file, pdf_file)

    # Convert PDF pages to images

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
  # Replace with your desired page numbers
    images = convert_from_path(pdf_file)

    for page_number, image in enumerate(images, start=1):
        if page_number in specific_pages:
            image_filename = os.path.join(output_folder, f"page_{page_number}.jpg")
            image.save(image_filename, "JPEG")
            print(f"Page {page_number} converted to image: {image_filename}")

    # Clean up the temporary PDF file
    os.remove(pdf_file)




def count_sections(doc):

    try:
        num_sections = len(doc.sections)
        return num_sections
    except Exception as e:
        print("Error:", e)
        return None

    return section_count
def list_images_in_folder(folder_path=r'static\images'):
    image_extensions = [".jpg", ".jpeg", ".png", ".gif", ".bmp"]  # Add more extensions as needed
    image_files = []

    for file in os.listdir(folder_path):
        if any(file.lower().endswith(ext) for ext in image_extensions):
            image_files.append('../static/images/'+file)

    return image_files


import os
from docx2pdf import convert
from pdf2image import convert_from_path
import pythoncom

def docx_to_pdf(input_docx):
    convert(input_docx, r'temp.pdf')


def pdf_to_images():
    input_pdf = r"temp.pdf"
    output_folder = r"static\images"
    images = convert_from_path(input_pdf)
    os.makedirs(output_folder, exist_ok=True)

    for i, image in enumerate(images):
        image.save(f"{output_folder}/page_{i + 1}.png", "PNG")

def add_arrow_to_word(doc):
    image_folder = r"static\temp_images"
    # Iterate through all image files in the folder
    count = 0
    for filename in os.listdir(image_folder):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            image_path = os.path.join(image_folder, filename)
            if count % 4 == 0:
                p = doc.add_paragraph()
            elif count ==0:
                p = doc.add_paragraph()
            count += 1
            r = p.add_run()
            r.add_picture(image_path,width=Inches(2))

    doc.add_section(WD_SECTION.NEW_PAGE)
    # Save the document
    doc.save('test_template.docx')
def main(input_docx):
    pythoncom.CoInitialize()
    docx_to_pdf(input_docx)
    pdf_to_images()

    os.remove('temp.pdf')  # Clean up the temporary PDF file


def get_image_links():
    folder_path = r"static\images"
    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp']
    image_file_info = []  # List to store (file_path, modification_time) tuples

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            _, ext = os.path.splitext(file)
            if ext.lower() in image_extensions:
                file_path = os.path.join(root, file)
                modification_time = os.path.getmtime(file_path)
                image_file_info.append((file_path, modification_time))

    # Sort the image file info list based on modification time
    sorted_image_file_info = sorted(image_file_info, key=lambda x: x[1])

    # Extract image links from sorted file info list
    image_links = ['../static/images/' + os.path.basename(file) for file, _ in sorted_image_file_info]

    return image_links
def create_temp_document(doc):
    # dimensional change
    image_pages = []
    image_pages.append(count_sections(doc)-1)
    print(image_pages)
    from Dimensional_stability_to_washing import dimensional_test_data, dimensional_test_template
    test_name, test_method, remarks, data = dimensional_test_data()
    input_data_list = [["±5", "±7", "±8", "±10"]]
    input_data_list = requirement_data(input_data_list)
    dimensional_test_template(doc, test_name, test_method, remarks, data, '+-5')
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()

    # appearance after washing
    image_pages.append(count_sections(doc)-1)
    from Appearance_after_washing import appearance_test_data, appearance_test_template
    test_name, test_method, remarks, title, req, sample_results = appearance_test_data()
    appearance_test_template(doc, test_name, test_method, remarks, title, ['Class 3-4', 'Class 4-5', 'Satisfactory', 'Class 4-5', 'Not Accepted'], sample_results)
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()

    # colour fastness to water
    image_pages.append(count_sections(doc)-1)
    from Colour_fastness_to_water import color_test_data, color_test_template
    test_name, test_method, remarks, title, values = color_test_data()
    color_test_template(doc, test_name, test_method, remarks, title, values,['Change 3-4','/','Staining 3-4'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()

    # colour fastness to perspiration
    image_pages.append(count_sections(doc)-1)
    from Colour_fastness_to_perspiration import color_fp_test_data, color_fp_test_template
    test_name, test_method, remarks, title, values = color_fp_test_data()
    color_fp_test_template(doc, test_name, test_method, remarks, title, values,['Change 3-4','Contrast staining 4-5','Staining 3-4'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()

    # colour fastness to water
    image_pages.append(count_sections(doc)-1)
    from Colour_fastness_to_rubbing import Color_fr_test_data, Color_fr_test_template
    test_name, test_method, remarks, title, values = Color_fr_test_data()
    Color_fr_test_template(doc, test_name, test_method, remarks, title, values,['Dry 4', 'Wet 2-3'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    sample_heading = doc.add_paragraph()
    sample_run = sample_heading.add_run("\n")

    # fiber composition
    image_pages.append(count_sections(doc)-1)
    from fiber_composition import fiber_test_template, fiber_test_data
    test_name = "Fiber Composition:"
    test_method = "DIN EN ISO 1833-7:2017,11:2017"
    remarks = ''
    test_name, test_method, remarks, title, sample_results = fiber_test_data()
    fiber_test_template(doc, test_name, test_method, remarks, title, sample_results,['±3%'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Alkylphenole test
    image_pages.append(count_sections(doc)-1)
    from Alkylphenole import Alkylphenole_test_data, Alkylphenole_test_template
    test_name, test_method, remarks, title, cas_no, values = Alkylphenole_test_data()
    Alkylphenole_test_template(doc, test_name, test_method, remarks, title, cas_no, values,['Sum of \nNP, OP: 5 mg/kg', 'Sum of \nNPEO, OPEO: 50 mg/kg','/'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Fabric Weight
    image_pages.append(count_sections(doc)-1)
    from fabric_weight import fabric_test_data, fabric_test_template
    test_name, test_method, remarks, title, values = fabric_test_data()
    fabric_test_template(doc, test_name, test_method, remarks, title, values,['±5%', '190 g/m²', '5.60 Oz'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Seam spirality after laundering:
    image_pages.append(count_sections(doc)-1)
    from Seam_spirality_after_laundering import seam_test_template, seam_test_data
    test_name, test_method, remarks, title, values = seam_test_data()
    seam_test_template(doc, test_name, test_method, remarks, title, values,['±3%', '3 cm'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Azo-Dyes (including Aniline):
    image_pages.append(count_sections(doc)-1)
    from Azo_Dyes import Azo_test_template, Azo_test_data
    test_name, test_method, remarks, values = Azo_test_data()
    Azo_test_template(doc, test_name, test_method, remarks, values,['20'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Aromatic Amine Salts:
    image_pages.append(count_sections(doc)-1)
    from Aromatic_Amine_Salts import Aromatic_test_data, Aromatic_test_template
    test_name, test_method, remarks, values = Aromatic_test_data()
    Aromatic_test_template(doc, test_name, test_method, remarks, values,['20'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Formaldehyde
    image_pages.append(count_sections(doc)-1)
    from Formaldehyde import Formaldehyde_test_data, Formaldehyde_test_template
    test_name, test_method, remarks, values = Formaldehyde_test_data()
    Formaldehyde_test_template(doc, test_name, test_method, remarks, values,['ND'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Total Lead
    image_pages.append(count_sections(doc)-1)
    from Total_Lead_Content import Total_test_data, Total_test_template
    test_name, test_method, remarks, values = Total_test_data()
    Total_test_template(doc, test_name, test_method, remarks, values,['Individual Test: 500 \n2-1 composite: 225\n3-1 composite: 150','75','90'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    # Total Cadmium (Cd) Content:
    image_pages.append(count_sections(doc)-1)
    from Total_Cadmium import Cadmium_test_template, Cadmium_test_data
    test_name, test_method, remarks, values = Cadmium_test_data()
    Cadmium_test_template(doc, test_name, test_method, remarks, values,['Individual Test: 100 \n2-1 composite: 45\n3-1 composite: 30','40'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    # Extractable Metals
    image_pages.append(count_sections(doc)-1)
    from Extractable_metals import extract_test_data, extract_test_template
    test_name, test_method, remarks, title, cas_no, mdl, values = extract_test_data()
    extract_test_template(doc, test_name, test_method, remarks, title, cas_no, mdl, values,['0.2','0.2','0.1','0.5','30','1.0','1.0','25','1.0','0.02','1000','100'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Quinoline
    image_pages.append(count_sections(doc)-1)
    from Quinoline import Quinoline_test_template, Quinoline_test_data
    test_name, test_method, remarks, values = Quinoline_test_data()
    Quinoline_test_template(doc, test_name, test_method, remarks, values,['Individual Test :50 \n2-1 composite: 22.5\n3-1 composite: 15','50'])
    doc.add_section(WD_SECTION.NEW_PAGE)

    # Phthalates
    image_pages.append(count_sections(doc)-1)
    from phthalates import Ph_test_data, Ph_test_template
    test_name, test_method, remarks, values = Ph_test_data()
    Ph_test_template(doc, test_name, test_method, remarks, values,['Individual Test :1000\n2-1 composite: 450\n3-1 composite: 300','Each :100 \nSum of all: 250'])
    doc.add_section(WD_SECTION.NEW_PAGE)
    # Polycyclic
    image_pages.append(count_sections(doc)-1)
    from Polycyclic import Poly_test_data, Poly_test_template
    test_name, test_method, remarks, title, cas_no, values = Poly_test_data()
    Poly_test_template(doc, test_name, test_method, remarks, title, cas_no, values,['8 PAHs \n(substance limit):\n0.5 mg/kg','2.0 mg/kg','no single limit','5.0 mg/kg'])
    from header import set_header
    set_header(doc)
    doc.save("temp_template.docx")
    return image_pages


def add_row_to_table(doc,number, test_name,result):

    # Get the table by index
    tables = doc.tables

    table = tables[4]

    # Add a new row to the table
    row = table.add_row()
    row_cells = row.cells
    row_cells[0].text = str(number-1)
    row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    row_cells[1].text = test_name
    if result == 'Pass':
        row_cells[2].text = 'X'
        row_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    elif result == 'Fail':
        row_cells[3].text = 'X'
        row_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    else:
        row_cells[4].text = 'See Actual Result'
        row_cells[4].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    row.height = Pt(20)

from docx import Document
from docx.shared import Inches  # Import Inches for image sizing

from docx import Document
from docx.shared import Inches

def add_care_labels(doc, labels):
    # Get the table by index
    folder_path = os.getcwd()
    for key,value in labels.items():
        if value:
            value.replace('\\','')
            labels[key] = folder_path+value[2:len(value)]
    tables = doc.tables
    table = tables[2]
    cell = table.cell(0, 1)

    # Add a new row to the table
    care_table = cell.add_table(rows=1, cols=5)
    care_table.autofit = False  # Prevent the table from automatically adjusting cell widths
    # Get the first row of the newly added table
    care_table_row = care_table.rows[0]
    index = 0
    # Iterate through the labels dictionary and add cells with images and text
    for section, image_path in labels.items():
        cell = care_table_row.cells[index]
        index += 1

        # Add the image and section name to the cell
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(0.4))  # You can adjust the width
        paragraph.alignment = 0  # Center-align the content
    widths = (Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4), Inches(0.4))
    set_col_widths(care_table, widths)
    doc.save('test_template.docx')


def add_arrow(label_name, flip):
    img = Image.open(r'static\arrow.png')
    image_width, image_height = img.size
    transparent_section_width = 200  # Adjust this value as needed

    new_img = Image.new('RGBA', (image_width + transparent_section_width, image_height), (0, 0, 0, 0))

    if flip:
        img = img.rotate(180)  # Rotate the arrow image 180 degrees

    if flip:
        # When flip is True, paste the arrow image at the right side of the new image
        new_img.paste(img, (0, 0))
    else:
        # When flip is False, paste the arrow image at the transparent section
        new_img.paste(img, (transparent_section_width, 0))

    draw = ImageDraw.Draw(new_img)

    font_path = r'static\cambria.ttf'
    font_size = 40
    myFont = ImageFont.truetype(font_path, font_size)

    label_text = label_name
    if flip:
        label_position = (image_width-15, 25)  # Adjust position for right corner
    else:
        label_position = (170, 25)  # Adjust position for left corner

    label_color = (232, 93, 68)  # RGB color for #e85d44
    draw.text(label_position, label_text, font=myFont, fill=label_color)

    edited_image_path = fr"static\temp_images\{label_name}.png"
    new_img.save(edited_image_path)
    return edited_image_path

from docx import Document
from PIL import Image

def insert_image_after_page(doc=1, image_path=r"static\edited_image.png", after_page=4, save_as='test_template.docx'):
    """
    Inserts an image after a specified page in a Word document.

    Parameters:
    - doc (Document): The loaded Word document object.
    - image_path (str): Path to the image to be inserted.
    - after_page (int): Page number after which to insert the image. Default is 4.
    - save_as (str): Path to save the modified document. Default is 'test_template.docx'.

    Returns:
    None
    """
    doc = Document('test_template.docx')
    # Try to find the end of the specified page (assuming each page is separated by a section break)
    try:
        section = doc.sections[after_page - 1]
        last_paragraph = section.footer.paragraphs[0]
        # Insert a page break after the last paragraph of the specified page
        last_paragraph.add_run().add_break()
    except IndexError:
        print("The document does not have the specified number of pages or does not use section breaks as delimiters.")
        return

    # Get the width of the original image in inches

    with Image.open(image_path) as img:

        original_width_in_inches = img.width / 96  # Assuming a DPI of 96 for the image

    # Add picture to the new page with its original width
    doc.add_picture(image_path, width=original_width_in_inches)

    # Save the modified document
    doc.save(save_as)

def set_primary_req(fabric,test_package,color):
    req_data = {'Appearance after washing':['Class 3-4', 'Class 4-5', 'Satisfactory', 'Class 4-5', 'Not Accepted'],'Fiber Composition:':['±3%'],'Colorfastness to saliva and perspiration':['Grade 5'],
                'Tensile test on small parts':'90N'}
    fabric_dictionary = {
        "woven": ["poplin", "twill", "denim","woven", "satin", "chiffon", "canvas", "flannel", "broadcloth", "gingham","herringbone", "jacquard", "seersucker", "organza", "gabardine", "brocade", "duck", "georgette","silk", "linen", "madras"],
        "knitted": ["jersey", "single jersey", "double jersey","knit","knitted", "rib knit", "1x1 rib", "2x2 rib", "purl knit","interlock", "french terry", "cable knit", "double knit", "ponte de roma", "milano rib","jacquard knit", "raschel knit", "tricot", "velour", "terry knit", "fleece"]}

    if fabric.lower() in fabric_dictionary['woven'] or 'jersey' in fabric.lower() or '1x1 rib' in fabric.lower() or 'viscose' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±5']
    elif 'terry' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±7']
    elif 'fine rib' in fabric.lower() or 'double rib' in fabric.lower() or 'interlock' in fabric.lower() or 'elastane' in fabric.lower() or 'viscose' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±8']
    elif 'muslin nappy' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±10']
    else:
        req_data['Dimensional stability to washing'] = ['±5']
    if fabric.lower() in fabric_dictionary['woven'] and 'P1' not in test_package:
        req_data['Seam spirality after laundering:'] = ['±3%' , '3cm']
    elif fabric.lower() in fabric_dictionary['knitted'] and 'P1' not in test_package:
        req_data['Seam spirality after laundering:'] = ['±5%' ,'5cm']
    else:
        req_data['Seam spirality after laundering:'] = ['up to 2cm', '2cm']
    if fabric.lower() in fabric_dictionary['woven']:
        req_data['Fabric weight:'] = ['±5%','190 g/m²','5.60 Oz']
    elif fabric.lower() in fabric_dictionary['woven']:
        req_data['Fabric weight:'] = ['±5%/+10%','190 g/m²','5.60 Oz']
    else:
        req_data['Fabric weight:'] = ['±5%','190 g/m²','5.60 Oz']
    if 'coruroy' in fabric.lower() or 'velvet' in fabric.lower() or 'velour' in fabric.lower() or 'flock print' in fabric.lower():
        req_data['Colour fastness to rubbing:'] = ['Dry 3', 'Wet 2']
    elif 'denim' in fabric.lower() or 'flannel' in fabric.lower() or 'peached' in fabric.lower() or 'pigment':
        req_data['Colour fastness to rubbing:'] = ['Dry 2-3', 'Wet 2']
    elif 'denim dark blue' in color.lower() or 'black denim' in color.lower() or 'overdyed' in color.lower() or 'special wash' in color.lower():
        req_data['Colour fastness to rubbing:'] = ['no testing', 'no testing']
    else:
        req_data['Colour fastness to rubbing:'] = ['Dry 3-4', 'Wet 2-3']
    req_data['Colour fastness to water'] = ['Grade 3-4','Grade 4-5','Grade 3-4']
    req_data['Colour fastness to perspiration'] = ['Change 3-4','/','Staining 3-4']
    return req_data

def set_secondarya_req(fabric,test_package,color):
    req_data = {'Appearance after washing':['Class 3-4', 'Class 4-5', 'Satisfactory', 'Class 4-5', 'Not Accepted'],'Fiber Composition:':['±3%'],'Colorfastness to saliva and perspiration':['Grade 5'],
                'Tensile test on small parts':'90N'}
    fabric_dictionary = {
        "woven": ["poplin", "twill", "denim","woven", "satin", "chiffon", "canvas", "flannel", "broadcloth", "gingham","herringbone", "jacquard", "seersucker", "organza", "gabardine", "brocade", "duck", "georgette","silk", "linen", "madras"],
        "knitted": ["jersey", "single jersey", "double jersey","knit","knitted", "rib knit", "1x1 rib", "2x2 rib", "purl knit","interlock", "french terry", "cable knit", "double knit", "ponte de roma", "milano rib","jacquard knit", "raschel knit", "tricot", "velour", "terry knit", "fleece"]}

    if fabric.lower() in fabric_dictionary['woven'] or 'jersey' in fabric.lower() or '1x1 rib' in fabric.lower() or 'viscose' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±5']
    elif 'terry' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±7']
    elif 'fine rib' in fabric.lower() or 'double rib' in fabric.lower() or 'interlock' in fabric.lower() or 'elastane' in fabric.lower() or 'viscose' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±8']
    elif 'muslin nappy' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±10']
    else:
        req_data['Dimensional stability to washing'] = ['±5']
    if fabric.lower() in fabric_dictionary['woven'] and 'P1' not in test_package:
        req_data['Seam spirality after laundering:'] = ['±3%' , '3cm']
    elif fabric.lower() in fabric_dictionary['knitted'] and 'P1' not in test_package:
        req_data['Seam spirality after laundering:'] = ['±5%' ,'5cm']
    else:
        req_data['Seam spirality after laundering:'] = ['up to 2cm', '2cm']
    if fabric.lower() in fabric_dictionary['woven']:
        req_data['Fabric weight:'] = ['±5%','190 g/m²','5.60 Oz']
    elif fabric.lower() in fabric_dictionary['woven']:
        req_data['Fabric weight:'] = ['±5%/+10%','190 g/m²','5.60 Oz']
    else:
        req_data['Fabric weight:'] = ['±5%','190 g/m²','5.60 Oz']
    if 'coruroy' in fabric.lower() or 'velvet' in fabric.lower() or 'velour' in fabric.lower() or 'flock print' in fabric.lower():
        req_data['Colour fastness to rubbing:'] = ['Dry 3', 'Wet 2']
    elif 'denim' in fabric.lower() or 'flannel' in fabric.lower() or 'peached' in fabric.lower() or 'pigment':
        req_data['Colour fastness to rubbing:'] = ['Dry 2-3', 'Wet 2']
    elif 'denim dark blue' in color.lower() or 'black denim' in color.lower() or 'overdyed' in color.lower() or 'special wash' in color.lower():
        req_data['Colour fastness to rubbing:'] = ['no testing', 'no testing']
    else:
        req_data['Colour fastness to rubbing:'] = ['Dry 3-4', 'Wet 2-3']
    req_data['Colour fastness to water'] = ['Grade 3-4','Grade 4-5','Grade 3-4']
    req_data['Colour fastness to perspiration'] = ['Change 3-4','/','Staining 3-4']
    return req_data


def mainPageRemainingData():
    # Initialize an empty dictionary to store key-value pairs
    result_dict = {}

    # Load the DOCX file
    doc = Document('component_format.docx')

    # Check if the document has tables
    if doc.tables:
        # Get the first table in the document
        table = doc.tables[0]

        # Iterate through the rows in the table
        for row in table.rows:
            # Ensure there are at least two cells in the row
            if len(row.cells) >= 2:
                # Extract the text from the first and second columns
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()

                # Add the key-value pair to the result_dict
                result_dict[key] = value

    return result_dict



def insert_image_after_page_5(docx_file='test_template.docx'):
    # Load the existing Word document
    doc = Document(docx_file)
    image_file = r'static\edited_image.png'
    # Get a reference to the 6th page (Python uses 0-based indexing)
    sixth_page = doc.sections[5]

    # Create a new page break after the 6th page
    doc.add_section(sixth_page.start_type)
    new_section = doc.sections[-1]
    new_section.start_type = sixth_page.start_type
    new_section.start_new_page = True

    # Insert the image on the new page
    image = Image.open(image_file)
    width, height = image.size
    image_width = Inches(6)  # You can adjust this as needed
    image_height = image_width * (height / width)

    doc.add_picture(image_file, width=image_width, height=image_height)

    # Save the modified document
    doc.save(docx_file)


def set_secondaryb_req(fabric,test_package,color):
    req_data = {'Appearance after washing':['Class 3-4', 'Class 4-5', 'Satisfactory', 'Class 4-5', 'Not Accepted'],'Fiber Composition:':['±3%'],'Colorfastness to saliva and perspiration':['Grade 5'],
                'Tensile test on small parts':'90N'}
    fabric_dictionary = {
        "woven": ["poplin", "twill", "denim","woven", "satin", "chiffon", "canvas", "flannel", "broadcloth", "gingham","herringbone", "jacquard", "seersucker", "organza", "gabardine", "brocade", "duck", "georgette","silk", "linen", "madras"],
        "knitted": ["jersey", "single jersey", "double jersey","knit","knitted", "rib knit", "1x1 rib", "2x2 rib", "purl knit","interlock", "french terry", "cable knit", "double knit", "ponte de roma", "milano rib","jacquard knit", "raschel knit", "tricot", "velour", "terry knit", "fleece"]}

    if fabric.lower() in fabric_dictionary['woven'] or 'jersey' in fabric.lower() or '1x1 rib' in fabric.lower() or 'viscose' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±5']
    elif 'terry' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±7']
    elif 'fine rib' in fabric.lower() or 'double rib' in fabric.lower() or 'interlock' in fabric.lower() or 'elastane' in fabric.lower() or 'viscose' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±8']
    elif 'muslin nappy' in fabric.lower():
        req_data['Dimensional stability to washing'] = ['±10']
    else:
        req_data['Dimensional stability to washing'] = ['±5']
    if fabric.lower() in fabric_dictionary['woven'] and 'P1' not in test_package:
        req_data['Seam spirality after laundering:'] = ['±3%' , '3cm']
    elif fabric.lower() in fabric_dictionary['knitted'] and 'P1' not in test_package:
        req_data['Seam spirality after laundering:'] = ['±5%' ,'5cm']
    else:
        req_data['Seam spirality after laundering:'] = ['up to 2cm', '2cm']
    if fabric.lower() in fabric_dictionary['woven']:
        req_data['Fabric weight:'] = ['±5%','190 g/m²','5.60 Oz']
    elif fabric.lower() in fabric_dictionary['woven']:
        req_data['Fabric weight:'] = ['±5%/+10%','190 g/m²','5.60 Oz']
    else:
        req_data['Fabric weight:'] = ['±5%','190 g/m²','5.60 Oz']
    if 'coruroy' in fabric.lower() or 'velvet' in fabric.lower() or 'velour' in fabric.lower() or 'flock print' in fabric.lower():
        req_data['Colour fastness to rubbing:'] = ['Dry 3', 'Wet 2']
    elif 'denim' in fabric.lower() or 'flannel' in fabric.lower() or 'peached' in fabric.lower() or 'pigment':
        req_data['Colour fastness to rubbing:'] = ['Dry 2-3', 'Wet 2']
    elif 'denim dark blue' in color.lower() or 'black denim' in color.lower() or 'overdyed' in color.lower() or 'special wash' in color.lower():
        req_data['Colour fastness to rubbing:'] = ['no testing', 'no testing']
    else:
        req_data['Colour fastness to rubbing:'] = ['Dry 3-4', 'Wet 2-3']
    req_data['Colour fastness to water'] = ['Grade 3-4','Grade 4-5','Grade 3-4']
    req_data['Colour fastness to perspiration'] = ['Change 3-4','/','Staining 3-4']
    return req_data