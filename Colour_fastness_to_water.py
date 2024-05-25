from docx import Document
from docx.shared import Pt,Inches
from global_func import set_col_widths


from create_test import create_test
def color_test_template(doc,test_name,test_method,remarks,title,result,req):
    create_test(doc, test_name, test_method, remarks)
    table = doc.add_table(rows=10, cols=len(result) + 1)
    table.style = 'Table Grid'
    table.autofit = True
    headers = ['Sample']
    for sample in result:
        temp = list(sample.keys())
        temp = temp[0]
        headers.append(temp[-1])
    for j, header in enumerate(headers):
        row = table.rows[0]
        Nombre_text_formatted = row.cells[j].paragraphs[0].add_run(header)
        Nombre_text_formatted.bold = True
        Nombre_text_formatted.font.size = Pt(12)
    for i, row in enumerate(table.rows):
        if i > 0:
            cell = row.cells[0]
            cell.text = title.pop(0)
    for index, res in enumerate(result, start=1):
        for i, row in enumerate(table.rows):
            if i > 0:
                cell = row.cells[index]
                temp = list(result[index - 1].values())
                temp = temp[0]
                cell.text = temp[i - 1]
    table.add_column(Pt(2))
    a = table.cell(1, len(result) + 1)
    a.text = req[0]
    a = table.cell(2, len(result) + 1)
    a.text = req[1]
    a = table.cell(3, len(result) + 1)
    for i in range(3, len(result) + 1):
        b = table.cell(i + 1, len(result) + 1)
        A = a.merge(b)
        a = A
    a.text = req[2]
    widths = (Inches(2), Inches(1),Inches(1),Inches(1),Inches(1),Inches(1),Inches(1),Inches(1),Inches(1))
    set_col_widths(table, widths)
    for row in table.rows:
        row.cells[-1].width = Inches(2)

def color_test_data():
    title = ['Color Change','Self-staining','Staining On:',' - Acetate',' - Cotton'	,' - Polyamide',' - Polyester',' - Acrylic	',' - Wool']
    values = [{'Sample A': ['4-5', '4-5', '-', '4-5', '4-5', '4-5', '4-5', '4-5', '4-5']}
        , {'Sample B': ['4-5', '4-5', '-', '4-5', '4-5', '4-5', '4-5', '4-5', '4-5']},
              {'Sample C': ['4-5', '4-5', '-', '4-5', '4-5', '4-5', '4-5', '4-5', '4-5']}
        , {'Sample D': ['4-5', '4-5', '-', '4-5', '4-5', '4-5', '4-5', '4-5', '4-5']},
              {'Sample E': ['4-5', '4-5', '-', '4-5', '4-5', '4-5', '4-5', '4-5', '4-5']}
        , {'Sample F': ['4-5', '4-5', '-', '4-5', '4-5', '4-5', '4-5', '4-5', '4-5']},
              {'Sample G': ['4-5', '4-5', '-', '4-5', '4-5', '4-5', '4-5', '4-5', '4-5']}
        , {'Sample H': ['4-5', '4-5', '-', '4-5', '4-5', '4-5', '4-5', '4-5', '4-5']}
              ]
    test_name = "Colour fastness to water:"
    test_method = "DIN EN ISO 105-E04"
    remarks = ''
    return test_name,test_method,remarks,title,values




