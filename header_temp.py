from docx import Document
from docx.shared import Pt,Inches
from global_func import set_col_widths
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT,WD_TABLE_ALIGNMENT
def set_header(doc):
    sections = doc.sections

    total_pages = len(doc.sections)
    recent_page = 1
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    header = section.header

    tti_logo = r"Logo\tti_logo.png"
    kik_logo = r"Logo\kik_logo.png"
    test_logo = r"Logo\test_logo.png"
    logos = [tti_logo,kik_logo,test_logo]
    table = header.add_table(5, 4,Inches(7))
    table.style = 'Table Grid'
    #table.rows[0].style = "borderColor:red;background-color:gray"
    header_set = ['Report No:','Receiving Date:','Issue Date:','Lab Location:','Page']
    header_input = {"Report No:":'02249-23','Receiving Date:':'Feb 02, 2023','Issue Date:':'Feb 06, 2023','Lab Location:':'Lahore (Pakistan)','Page':f'{recent_page} of {total_pages}'}
    a = table.cell(0, 1)
    for j in range(0, 5):
        b = table.cell(j, 1)
        A = a.merge(b)
        a = A
    table1 = a.add_table(5,2)
    table1.style = "Table Grid"
    for i in range(5):
        row = table1.rows[i]
        row_cells = row.cells[1].paragraphs[0].add_run(header_set[i])
        row_cells.bold = True
        row_cells.font.size = Pt(10)
        row_cells1 = row.cells[2].paragraphs[0].add_run(header_input[header_set[i]])
        row_cells1.bold = True
        row_cells1.font.size = Pt(10)

    for i in [0,2,3]:
        a = table.cell(0,i)
        for j in range(0,5):
            b = table.cell(j,i)
            A = a.merge(b)
            a = A
        paragraph = a.paragraphs[0]
        run = paragraph.add_run()
        if i ==0:
            run.add_picture(logos[0], width=1000000, height=1000000)
        else:
            run.add_picture(logos[i-2], width=1000000, height=1000000)
    table.autofit = False
    table.allow_autofit = False
    widths = (Inches(1.4), Inches(1.6), Inches(1.5), Inches(1.4),Inches(1.25))
    set_col_widths(table,widths)
    for row in table.rows:
        row.height = Inches(0.2)
